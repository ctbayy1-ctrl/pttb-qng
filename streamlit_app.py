import streamlit as st
import pandas as pd
from datetime import datetime
import time
import os
import io
import xml.etree.ElementTree as ET
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
from selenium.common.exceptions import WebDriverException, TimeoutException, NoSuchElementException
import json
import httpx
import asyncio

# --- CẤU HÌNH CÁC CHỈ TIÊU XML ---
XML_TAG_MAP = {
    '01/GTGT': {
        'doanh_thu_kct': 'ct26', 'doanh_thu_0': 'ct29', 'doanh_thu_5': 'ct30',
        'doanh_thu_10': 'ct32', 'thue_gtgt_duoc_khau_tru': 'ct25', 'thue_gtgt_phai_nop': 'ct40',
        'tong_doanh_thu_thue': 'ct34'
    },
    '03/TNDN': {
        'tong_doanh_thu_pl': 'ct04',
        'thu_nhap_khac': 'ct19',
        'loi_nhuan_truoc_thue': 'ctA1',
        'chi_phi': 'ct12',
        'thue_tndn_phai_nop': 'ctC9'
    },
    'BCTC_CDKT': {
        'tien': 'ct110', 'hang_ton_kho': 'ct140',
    },
    'BCTC_KQKD': {
        'doanh_thu_bh': 'ct01', 'gia_von': 'ct11', 'chi_phi_lai_vay': 'ct23',
        'loi_nhuan_truoc_thue': 'ct50',
    },
    '05/KK-TNCN': {
        'tong_so_ld': 'ct16',
        'tong_tnct': 'ct21',
        'tong_thue_da_khau_tru': 'ct29'
    },
    '05/QTT-TNCN': {
        'tong_so_ld': 'ct16',
        'tong_tnct': 'ct23',
        'tong_thue_da_khau_tru': 'ct31',
        'pl_ho_ten': 'ct07',
        'pl_mst': 'ct08',
        'pl_tnct': 'ct12',
        'pl_thue_da_khau_tru': 'ct22'
    }
}
# Ánh xạ mã tờ khai sang loại để xử lý
MA_TK_MAP = {
    '842': '01/GTGT', '844': '02/GTGT', '950': '03/TNDN', '892': '03/TNDN',
    '402': 'BCTC', '699': 'BCTC',
    '683': 'BCTC',
    '953': '05/QTT-TNCN',
    '864': '05/KK-TNCN',
}


# --- MODULE 4: GIAO DIỆN NGƯỜI DÙNG (USER INTERFACE) ---
def setup_ui():
    st.set_page_config(page_title="Ứng dụng Hỗ trợ Rà soát HSKT", layout="wide")
    st.title("🤖 Ứng dụng Hỗ trợ Rà soát Hồ sơ Khai thuế")

    loai_to_khai_list = [
        "01/GTGT - TỜ KHAI THUẾ GIÁ TRỊ GIA TĂNG (TT80/2021)",
        "03/TNDN - Tờ khai quyết toán thuế TNDN (TT80/2021)",
        "TT200 - Bộ báo cáo tài chính",
        "05/QTT-TNCN - TỜ KHAI QUYẾT TOÁN THUẾ THU NHẬP CÁ NHÂN (TT80/2021)",
        "05/KK-TNCN - Tờ khai khấu trừ thuế TNCN (TT80/2021)",
    ]

    st.sidebar.header("Thông tin Rà soát")
    mode = st.sidebar.radio("Chọn chế độ làm việc:", ("Tải lên Thủ công", "Tự động hóa"))

    params = {"mode": mode}

    if mode == "Tự động hóa":
        st.sidebar.subheader("Thông tin Đăng nhập")
        params["username"] = st.sidebar.text_input("Tài khoản", "")
        params["password"] = st.sidebar.text_input("Mật khẩu", type="password")
        st.sidebar.markdown("---")
        params["mst"] = st.sidebar.text_input("Mã số thuế cần rà soát", "", max_chars=14)
        params["nam_tinh_thue"] = st.sidebar.number_input("Năm tính thuế", min_value=2010,
                                                          max_value=datetime.now().year, value=datetime.now().year)
        params["loai_tks"] = st.sidebar.multiselect("Chọn (các) loại tờ khai", loai_to_khai_list)
    else:
        st.sidebar.subheader("Tải lên Hồ sơ (.xml)")
        params["uploaded_files"] = st.sidebar.file_uploader(
            "Chọn các tệp hồ sơ (.xml)",
            accept_multiple_files=True, type=['xml']
        )

    st.sidebar.markdown("---")
    st.sidebar.subheader("Cấu hình Phân tích")
    params["accounting_standard"] = st.sidebar.selectbox(
        "Chế độ kế toán (BCTC):",
        ["Chưa chọn", "Thông tư 133", "Thông tư 200"],
        help="Chọn chế độ kế toán áp dụng cho Báo cáo tài chính để bật các quy tắc đối chiếu phù hợp."
    )

    st.sidebar.subheader("Hóa đơn Đầu ra")
    params["output_invoice_type"] = st.sidebar.selectbox(
        "Loại Bảng kê Bán ra:",
        ["Tổng hợp", "Chi tiết"]
    )
    params["output_invoice_file"] = st.sidebar.file_uploader(
        "Tải lên Bảng kê Hóa đơn Bán ra (.xlsx, .csv)",
        type=['xlsx', 'csv'],
        key="output_uploader"
    )

    st.sidebar.subheader("Hóa đơn Đầu vào")
    params["input_invoice_files"] = st.sidebar.file_uploader(
        "Tải lên Bảng kê Hóa đơn Mua vào (.xlsx, .csv)",
        type=['xlsx', 'csv'],
        key="input_uploader",
        accept_multiple_files=True
    )

    st.sidebar.subheader("Tài liệu khác")
    params["financial_notes_file"] = st.sidebar.file_uploader(
        "Tải lên Thuyết minh BCTC (.docx, .pdf, .xlsx)",
        type=['docx', 'pdf', 'xlsx', 'xls', 'doc']
    )

    st.sidebar.markdown("---")
    st.sidebar.subheader("Cấu hình AI")
    params["gemini_api_key"] = st.sidebar.text_input("Gemini API Key", type="password",
                                                     help="Nhập API Key của bạn để sử dụng chức năng phân tích của Gemini.")

    params["start_button"] = st.sidebar.button("🚀 Bắt đầu Phân tích")
    return params


# --- MODULE 1: THU THẬP DỮ LIỆU (DATA ACQUISITION) ---
class WebScraper:
    def __init__(self, params):
        self.params = params
        self.download_dir = os.path.join(os.getcwd(), "hoso_thue_auto")
        if not os.path.exists(self.download_dir): os.makedirs(self.download_dir)
        self.driver = self._initialize_driver()

    def _initialize_driver(self):
        try:
            edge_options = webdriver.EdgeOptions()
            edge_options.add_argument("--headless")
            edge_options.add_argument("--window-size=1920,1080")
            prefs = {"download.default_directory": self.download_dir}
            edge_options.add_experimental_option("prefs", prefs)
            service = EdgeService(executable_path='./msedgedriver.exe')
            driver = webdriver.Edge(service=service, options=edge_options)
            st.info("Đang sử dụng trình duyệt Microsoft Edge (chế độ chạy ẩn).")
            return driver
        except WebDriverException:
            st.warning("Không tìm thấy Edge/msedgedriver.exe. Đang thử với Chrome...")
            try:
                chrome_options = webdriver.ChromeOptions()
                chrome_options.add_argument("--headless")
                chrome_options.add_argument("--window-size=1920,1080")
                prefs = {"download.default_directory": self.download_dir}
                chrome_options.add_experimental_option("prefs", prefs)
                service = ChromeService(executable_path='./chromedriver.exe')
                driver = webdriver.Chrome(service=service, options=chrome_options)
                st.info("Đang sử dụng trình duyệt Google Chrome (chế độ chạy ẩn).")
                return driver
            except WebDriverException as e:
                st.error("Không thể khởi tạo trình duyệt. Vui lòng kiểm tra file driver.")
                raise e

    def login(self):
        st.write(f"1. Đang đăng nhập vào http://thuedientu.tct.vn/...")
        try:
            self.driver.get("http://thuedientu.tct.vn/")
            wait = WebDriverWait(self.driver, 10)
            wait.until(EC.presence_of_element_located((By.ID, "_userName"))).send_keys(self.params['username'])
            self.driver.find_element(By.ID, "password").send_keys(self.params['password'])
            self.driver.find_element(By.ID, "dangnhap").click()
            wait.until(EC.presence_of_element_located((By.XPATH, "//div[contains(text(), 'Tra cứu')]")))
            st.success("Đăng nhập thành công!")
            return True
        except (TimeoutException, NoSuchElementException) as e:
            st.error(f"Lỗi đăng nhập: {e}")
            return False

    def search_declarations(self, loai_tk):
        st.write(f"   - Đang tra cứu cho: **{loai_tk}**")
        try:
            wait = WebDriverWait(self.driver, 10)
            actions = ActionChains(self.driver)
            self.driver.switch_to.default_content()
            tra_cuu_menu = wait.until(EC.visibility_of_element_located(
                (By.XPATH, "//div[@class='text_memu_ngang' and contains(text(), 'Tra cứu')]")))
            actions.move_to_element(tra_cuu_menu).perform()
            time.sleep(2)
            tra_cuu_tk_link = wait.until(
                EC.element_to_be_clickable((By.XPATH, "//a[contains(text(), 'Tra cứu tờ khai')]")))
            actions.move_to_element(tra_cuu_tk_link).click().perform()
            wait.until(EC.frame_to_be_available_and_switch_to_it((By.NAME, "rframe")))
            self.driver.find_element(By.ID, "mst").clear()
            self.driver.find_element(By.ID, "mst").send_keys(self.params['mst'])
            Select(self.driver.find_element(By.ID, "maTKhai")).select_by_visible_text(loai_tk)
            self.driver.find_element(By.XPATH, "//input[@value='Tra cứu']").click()
            return True
        except (TimeoutException, NoSuchElementException) as e:
            st.warning(f"Lỗi trong quá trình tra cứu cho '{loai_tk}': {e}")
            return False

    def analyze_and_download(self, loai_tk):
        try:
            wait = WebDriverWait(self.driver, 15)
            results_table = wait.until(EC.presence_of_element_located((By.ID, "tbl_content_search")))
            rows = results_table.find_elements(By.TAG_NAME, "tr")[1:]
            all_results = [{'ky_tinh_thue': r.find_elements(By.TAG_NAME, "td")[3].text,
                            'loai': r.find_elements(By.TAG_NAME, "td")[4].text,
                            'lan_bs': r.find_elements(By.TAG_NAME, "td")[5].text,
                            'trang_thai': r.find_elements(By.TAG_NAME, "td")[8].text,
                            'download_link': r.find_elements(By.TAG_NAME, "td")[2].find_element(By.TAG_NAME, 'a')} for r
                           in rows if len(r.find_elements(By.TAG_NAME, "td")) > 8]
            results_in_year = [r for r in all_results if str(self.params['nam_tinh_thue']) in r['ky_tinh_thue']]
            hskts_to_download = {}
            for row in results_in_year:
                if 'TMS - Gói tin hạch toán thành công' not in row['trang_thai']: continue
                ky = row['ky_tinh_thue']
                if row['loai'] == 'Bổ sung':
                    lan_bs_hien_tai = int(row['lan_bs']) if row['lan_bs'].isdigit() else 0
                    if ky not in hskts_to_download or hskts_to_download[ky][
                        'loai'] != 'Bổ sung' or lan_bs_hien_tai > int(hskts_to_download[ky].get('lan_bs', '0')):
                        hskts_to_download[ky] = row
                elif row['loai'] == 'Chính thức' and ky not in hskts_to_download:
                    hskts_to_download[ky] = row
            downloaded_files_for_tk = []
            for ky, hsk_info in hskts_to_download.items():
                st.write(f"     -> Đang tải về HSKT cho kỳ `{ky}`...")
                hsk_info['download_link'].click()
                time.sleep(3)
                latest_file = max([os.path.join(self.download_dir, f) for f in os.listdir(self.download_dir)],
                                  key=os.path.getctime)
                loai_tk_filename = loai_tk.split(' - ')[0].replace('/', '-')
                ky_filename = ky.replace('/', '-')
                new_filename = f"{self.params['mst']}_{loai_tk_filename}_{ky_filename}.xml"
                new_filepath = os.path.join(self.download_dir, new_filename)
                os.rename(latest_file, new_filepath)
                downloaded_files_for_tk.append(new_filepath)
            st.write(f"   -> Đã tải về {len(downloaded_files_for_tk)} hồ sơ cho loại tờ khai này.")
            return downloaded_files_for_tk
        except (TimeoutException, NoSuchElementException):
            st.warning(f"Không tìm thấy kết quả cho '{loai_tk}'.")
            return []
        finally:
            self.driver.switch_to.default_content()

    def close_driver(self):
        if hasattr(self, 'driver') and self.driver: self.driver.quit()


def process_summary_invoice_data(uploaded_file):
    if not uploaded_file:
        return None
    try:
        st.write("Đang xử lý file Bảng kê hóa đơn tổng hợp...")
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, skiprows=5)
        else:
            df = pd.read_excel(uploaded_file, skiprows=5)

        df.columns = [
            'STT', 'KyHieuMauSo', 'KyHieuHoaDon', 'SoHoaDon', 'NgayLap', 'MSTNguoiBan', 'TenNguoiBan',
            'MSTNguoiMua', 'TenNguoiMua', 'DiaChiNguoiMua', 'MaSoThueToChucCungCap', 'MaSoThueToChucTruyenNhan',
            'TongTienChuaThue', 'TongTienThue', 'TongTienChietKhau', 'TongTienPhi', 'TongTienThanhToan',
            'DonViTienTe', 'TyGia', 'TrangThaiHoaDon', 'KetQuaKiemTra'
        ]

        numeric_cols = ['TongTienChuaThue', 'TongTienThue', 'TongTienChietKhau', 'TongTienThanhToan', 'TyGia']
        for col in numeric_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

        statuses_to_sum = [
            'Hóa đơn mới',
            'Hóa đơn thay thế',
            'Hóa đơn điều chỉnh',
            'Hóa đơn đã bị điều chỉnh'
        ]

        invoices_for_summing = df[df['TrangThaiHoaDon'].isin(statuses_to_sum)].copy()

        for col in ['TongTienChuaThue', 'TongTienThue', 'TongTienChietKhau', 'TongTienThanhToan']:
            invoices_for_summing.loc[:, col] = invoices_for_summing.apply(
                lambda row: row[col] * row['TyGia'] if row['DonViTienTe'] != 'VND' else row[col],
                axis=1
            )

        total_pre_tax = invoices_for_summing['TongTienChuaThue'].sum()
        total_tax = invoices_for_summing['TongTienThue'].sum()
        total_discount = invoices_for_summing['TongTienChietKhau'].sum()
        total_payment = invoices_for_summing['TongTienThanhToan'].sum()

        summary = {
            'total_pre_tax': total_pre_tax,
            'total_tax': total_tax,
            'total_discount': total_discount,
            'total_payment': total_payment
        }

        st.success("Xử lý Bảng kê hóa đơn tổng hợp hoàn tất!")
        return {
            "valid_summary": summary,
            "full_df": df
        }
    except Exception as e:
        st.error(f"Lỗi khi xử lý file hóa đơn tổng hợp: {e}")
        return None


def process_detailed_invoice_data(uploaded_file):
    if not uploaded_file:
        return None
    try:
        st.write("Đang xử lý file Bảng kê hóa đơn chi tiết...")
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)

        df.columns = [
            'KyHieu', 'SoHoaDon', 'NgayLap', 'TenNguoiMua', 'MSTNguoiMua', 'TinhChat',
            'TenHangHoa', 'DonViTinh', 'SoLuong', 'DonGia', 'TienChietKhau', 'EmptyCol',
            'ThanhTien', 'ThueSuat', 'DonViTienTe', 'TyGia', 'TrangThaiHoaDon',
            'HD_LienQuan', 'Ngay_HD_LienQuan', 'DonViCungCap'
        ]
        df = df.drop(columns=['EmptyCol'])

        numeric_cols = ['SoLuong', 'DonGia', 'TienChietKhau', 'ThanhTien', 'ThueSuat', 'TyGia']
        for col in numeric_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        df['NgayLap'] = pd.to_datetime(df['NgayLap'], errors='coerce')

        df['ThanhTien_TinhToan'] = df['SoLuong'] * df['DonGia']
        mismatched_invoices = df[abs(df['ThanhTien'] - df['ThanhTien_TinhToan']) > 0.01].copy()
        mismatched_invoices = mismatched_invoices[
            ['SoHoaDon', 'NgayLap', 'TenHangHoa', 'SoLuong', 'DonGia', 'ThanhTien', 'ThanhTien_TinhToan']]

        statuses_to_sum = ['Hóa đơn mới', 'Hóa đơn thay thế', 'Hóa đơn điều chỉnh', 'Hóa đơn đã bị điều chỉnh']
        invoices_for_summing = df[df['TrangThaiHoaDon'].isin(statuses_to_sum)].copy()

        invoices_for_summing['TienChuaThue'] = invoices_for_summing['ThanhTien'] - invoices_for_summing['TienChietKhau']
        invoices_for_summing['TienThue'] = invoices_for_summing['TienChuaThue'] * invoices_for_summing['ThueSuat']

        for col in ['TienChuaThue', 'TienThue', 'TienChietKhau']:
            invoices_for_summing.loc[:, col] = invoices_for_summing.apply(
                lambda row: row[col] * row['TyGia'] if row['DonViTienTe'] != 'VND' else row[col],
                axis=1
            )

        total_pre_tax = invoices_for_summing['TienChuaThue'].sum()
        total_tax = invoices_for_summing['TienThue'].sum()
        total_discount = invoices_for_summing['TienChietKhau'].sum()
        total_payment = total_pre_tax + total_tax

        summary = {
            'total_pre_tax': total_pre_tax,
            'total_tax': total_tax,
            'total_discount': total_discount,
            'total_payment': total_payment
        }

        st.success("Xử lý Bảng kê hóa đơn chi tiết hoàn tất!")
        return {
            "valid_summary": summary,
            "mismatch_df": mismatched_invoices,
            "full_df": df
        }
    except Exception as e:
        st.error(f"Lỗi khi xử lý file hóa đơn chi tiết: {e}")
        return None


def process_input_invoice_data(uploaded_files):
    if not uploaded_files:
        return None

    all_dfs = []
    try:
        st.write(f"Đang xử lý {len(uploaded_files)} file Bảng kê hóa đơn đầu vào...")
        for uploaded_file in uploaded_files:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file, skiprows=5, header=None)
            else:
                df = pd.read_excel(uploaded_file, skiprows=5, header=None)

            if len(df.columns) >= 21:
                st.info(f"Đã nhận diện file '{uploaded_file.name}' là mẫu đầu vào loại 1 (data-1).")
                df.columns = [
                                 'STT', 'KyHieuMauSo', 'KyHieuHoaDon', 'SoHoaDon', 'NgayLap', 'MSTNguoiBan',
                                 'TenNguoiBan', 'DiaChiNguoiBan',
                                 'MSTNguoiMua', 'TenNguoiMua', 'CCCD', 'MaSothueCCGiaiPhap', 'MaSoThueTruyenNhan',
                                 'TongTienChuaThue', 'TongTienThue', 'TongTienChietKhau', 'TongTienThanhToan',
                                 'TrangThaiHoaDon', 'KetQuaKiemTra', 'Col1', 'Col2'
                             ][:len(df.columns)]
                df = df.drop(columns=['Col1', 'Col2', 'CCCD'], errors='ignore')
            elif len(df.columns) >= 19:
                st.info(f"Đã nhận diện file '{uploaded_file.name}' là mẫu đầu vào loại 2 (mtt1).")
                df.columns = [
                                 'STT', 'KyHieuMauSo', 'KyHieuHoaDon', 'SoHoaDon', 'NgayLap', 'MSTNguoiBan',
                                 'TenNguoiBan', 'DiaChiNguoiBan',
                                 'MSTNguoiMua', 'TenNguoiMua', 'CCCD', 'MaSothueCCGiaiPhap', 'MaSoThueTruyenNhan',
                                 'TongTienChuaThue', 'TongTienThue', 'TongTienChietKhau', 'TongTienThanhToan',
                                 'TrangThaiHoaDon', 'KetQuaKiemTra'
                             ][:len(df.columns)]
                df = df.drop(columns=['CCCD'], errors='ignore')
            else:
                st.warning(f"Không thể nhận diện mẫu cho file '{uploaded_file.name}'. Bỏ qua file này.")
                continue

            all_dfs.append(df)

        if not all_dfs:
            st.error("Không có file hóa đơn đầu vào hợp lệ nào được xử lý.")
            return None

        final_df = pd.concat(all_dfs, ignore_index=True)

        numeric_cols = ['TongTienChuaThue', 'TongTienThue', 'TongTienChietKhau', 'TongTienThanhToan']
        for col in numeric_cols:
            if col in final_df.columns:
                final_df[col] = pd.to_numeric(final_df[col], errors='coerce').fillna(0)

        if 'NgayLap' in final_df.columns:
            final_df['NgayLap'] = pd.to_datetime(final_df['NgayLap'], errors='coerce', dayfirst=True)

        total_pre_tax = final_df['TongTienChuaThue'].sum()
        total_tax = final_df['TongTienThue'].sum()
        total_discount = final_df['TongTienChietKhau'].sum()
        total_payment = final_df['TongTienThanhToan'].sum()

        summary = {
            'total_pre_tax': total_pre_tax,
            'total_tax': total_tax,
            'total_discount': total_discount,
            'total_payment': total_payment
        }

        st.success("Xử lý Bảng kê hóa đơn đầu vào hoàn tất!")
        return {
            "valid_summary": summary,
            "full_df": final_df
        }
    except Exception as e:
        st.error(f"Lỗi khi xử lý file hóa đơn đầu vào: {e}")
        return None


def process_financial_notes(uploaded_file):
    if not uploaded_file:
        return None
    st.write("Đang xử lý file Thuyết minh BCTC...")
    content = ""
    try:
        if uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                content += para.text + '\n'
            st.success("Trích xuất nội dung từ file .docx thành công!")
        elif uploaded_file.name.endswith('.pdf'):
            content = "Chức năng đọc file PDF sẽ được phát triển trong các phiên bản sau."
            st.info(content)
        elif uploaded_file.name.endswith(('.xlsx', '.xls')):
            content = "Chức năng đọc file Excel cho Thuyết minh BCTC sẽ được phát triển trong các phiên bản sau."
            st.info(content)
        else:
            content = "Định dạng file này chưa được hỗ trợ để trích xuất nội dung."
            st.warning(content)
        return content
    except Exception as e:
        st.error(f"Lỗi khi xử lý file Thuyết minh BCTC: {e}")
        return None


# --- MODULE 2 & 3: PHÂN TÍCH DỮ LIỆU & RỦI RO ---
def parse_and_analyze(files, accounting_standard, output_invoice_data, input_invoice_data, notes_content):
    st.write("Đang xử lý, bóc tách và tổng hợp dữ liệu...")
    all_declarations = []
    st.session_state['tndn_main_df'] = pd.DataFrame()
    st.session_state['tndn_appendix_df'] = pd.DataFrame()
    st.session_state['gtgt_detailed_df'] = pd.DataFrame()

    for file_path in files:
        data = parse_xml_data(file_path)
        # === GỠ LỖI: HIỂN THỊ DỮ LIỆU THÔ ĐÃ BÓC TÁCH ===
        st.subheader(f"Dữ liệu thô từ file: {os.path.basename(file_path)}")
        st.json(data)
        # ===============================================
        if not data: continue
        ma_tk = data.get('maTKhai')
        loai_tk_code = MA_TK_MAP.get(ma_tk)
        ky = data.get('kyKKhai') or "Không xác định"
        if loai_tk_code:
            all_declarations.append(
                {'loai_tk': loai_tk_code, 'ky': ky, 'data': data, 'filename': os.path.basename(file_path)})
            if loai_tk_code == '03/TNDN':
                st.session_state['tndn_main_df'] = generate_tndn_main_form_df(data)
                st.session_state['tndn_appendix_df'] = generate_tndn_appendix_03_1a_df(data)
            elif loai_tk_code == '01/GTGT':
                st.session_state['gtgt_detailed_df'] = generate_gtgt_detailed_df(all_declarations)

    if not all_declarations and not output_invoice_data and not input_invoice_data and not notes_content:
        st.warning("Không có dữ liệu nào được cung cấp để phân tích.")
        return [], [], pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()

    gtgt_summary_df = generate_gtgt_summary(all_declarations)
    gtgt_detailed_df = st.session_state.get('gtgt_detailed_df', pd.DataFrame())
    balance_sheet_df = generate_balance_sheet_df(all_declarations)
    income_statement_df = generate_income_statement_df(all_declarations)
    trial_balance_df = generate_trial_balance_df(all_declarations)
    tndn_summary_df = generate_tndn_summary(all_declarations)
    tncn_qtt_summary_df, tncn_details_df = generate_tncn_summary(all_declarations)
    tncn_kk_summary_df = generate_tncn_kk_summary(all_declarations)

    # === GỠ LỖI: KIỂM TRA CÁC DATAFRAME TRƯỚC KHI TRẢ VỀ ===
    st.markdown("---")
    st.subheader("Trạng thái Dữ liệu (DEBUG)")
    data_status = {
        "Tờ khai GTGT (Tổng hợp)": not gtgt_summary_df.empty,
        "Tờ khai GTGT (Chi tiết)": not gtgt_detailed_df.empty,
        "Bảng Cân đối Kế toán": not balance_sheet_df.empty,
        "Báo cáo KQKD": not income_statement_df.empty,
        "Bảng Cân đối Tài khoản": not trial_balance_df.empty,
        "Tờ khai TNDN (Tổng hợp)": not tndn_summary_df.empty,
        "Tờ khai TNCN QTT (Tổng hợp)": not tncn_qtt_summary_df.empty,
        "Tờ khai TNCN QTT (Chi tiết)": not tncn_details_df.empty,
        "Tờ khai TNCN KK (Tổng hợp)": not tncn_kk_summary_df.empty,
    }
    st.write(data_status)
    st.markdown("---")
    # =========================================================

    st.write("Đang phân tích rủi ro...")
    output_invoice_pre_tax_total = output_invoice_data.get('valid_summary', {}).get('total_pre_tax',
                                                                                    None) if output_invoice_data else None
    all_checks = run_risk_checks(all_declarations, gtgt_summary_df, tncn_kk_summary_df, accounting_standard,
                                 output_invoice_pre_tax_total, input_invoice_data)
    st.success("Phân tích hoàn tất!")
    return all_declarations, all_checks, gtgt_summary_df, gtgt_detailed_df, balance_sheet_df, income_statement_df, trial_balance_df, tndn_summary_df, tncn_qtt_summary_df, tncn_details_df, tncn_kk_summary_df


def parse_xml_data(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
        tree = ET.fromstring(xml_content)
        data = {}

        parent_map = {c: p for p in tree.iter() for c in p}
        single_value_tags = ['maTKhai', 'kyKKhai', 'mst', 'tenNNT', 'dchiNNT', 'tenHuyenNNT', 'tenTinhNNT', 'loaiTKhai',
                             'soLan']

        for elem in tree.iter():
            if not (elem.text and elem.text.strip()):
                continue

            tag = elem.tag.split('}', 1)[-1]
            key = tag
            value = elem.text.strip()

            parent = parent_map.get(elem)
            if parent is not None:
                parent_tag = parent.tag.split('}', 1)[-1]

                if parent_tag == 'NamNay' and tag.startswith('ct'):
                    key = f"kqkd_nn_{tag}"
                elif parent_tag == 'NamTruoc' and tag.startswith('ct'):
                    key = f"kqkd_nt_{tag}"
                elif parent_tag == 'SoCuoiNam' and tag.startswith('ct'):
                    key = f"scn_{tag}"
                elif parent_tag == 'SoDauNam' and tag.startswith('ct'):
                    key = f"sdn_{tag}"
                else:
                    grandparent = parent_map.get(parent)
                    if grandparent is not None:
                        grandparent_tag = grandparent.tag.split('}', 1)[-1]
                        if grandparent_tag == 'SoPhatSinhTrongKy':
                            if parent_tag == 'No':
                                key = f"ps_no_{tag}"
                            elif parent_tag == 'Co':
                                key = f"ps_co_{tag}"
                        elif grandparent_tag == 'SoDuDauKy':
                            if parent_tag == 'No':
                                key = f"sddk_no_{tag}"
                            elif parent_tag == 'Co':
                                key = f"sddk_co_{tag}"
                        elif grandparent_tag == 'SoDuCuoiKy':
                            if parent_tag == 'No':
                                key = f"sdck_no_{tag}"
                            elif parent_tag == 'Co':
                                key = f"sdck_co_{tag}"

            if tag in single_value_tags:
                if tag not in data: data[tag] = value
            else:
                if key in data:
                    if not isinstance(data[key], list): data[key] = [data[key]]
                    data[key].append(value)
                else:
                    data[key] = value
        return data
    except (ET.ParseError, FileNotFoundError) as e:
        st.error(f"Lỗi khi đọc tệp {os.path.basename(file_path)}: {e}")
        return {}


def get_single_value(data, key, default=0):
    value = data.get(key, default)
    try:
        if isinstance(value, list):
            return float(value[0]) if value else float(default)
        return float(value) if value is not None else float(default)
    except (ValueError, TypeError):
        return float(default)


def get_string_value(data, key, default=''):
    value = data.get(key, default)
    if isinstance(value, list):
        return str(value[0]) if value else str(default)
    return str(value) if value is not None else str(default)


def generate_gtgt_summary(declarations):
    gtgt_decls = [d for d in declarations if 'GTGT' in d['loai_tk']]
    if not gtgt_decls: return pd.DataFrame()
    periods = sorted(list(set(d['ky'] for d in gtgt_decls)))
    summary_data = {'Chỉ tiêu': ['Doanh thu chịu thuế GTGT', 'Thuế GTGT được khấu trừ', 'Thuế GTGT phải nộp']}
    for period in periods:
        p_data = next((d['data'] for d in gtgt_decls if d['ky'] == period), {})
        dt_total = sum(get_single_value(p_data, tag, 0) for tag in
                       [XML_TAG_MAP['01/GTGT']['doanh_thu_kct'], XML_TAG_MAP['01/GTGT']['doanh_thu_0'],
                        XML_TAG_MAP['01/GTGT']['doanh_thu_5'], XML_TAG_MAP['01/GTGT']['doanh_thu_10']])
        thue_kt = get_single_value(p_data, XML_TAG_MAP['01/GTGT']['thue_gtgt_duoc_khau_tru'], 0)
        thue_pn = get_single_value(p_data, XML_TAG_MAP['01/GTGT']['thue_gtgt_phai_nop'], 0)
        summary_data[period] = [dt_total, thue_kt, thue_pn]
    df = pd.DataFrame(summary_data)
    if len(periods) > 0: df['Tổng cộng'] = df[periods].sum(axis=1)
    return df


def generate_gtgt_detailed_df(declarations):
    gtgt_decl = next((d for d in reversed(declarations) if '01/GTGT' in d['loai_tk']), None)
    if not gtgt_decl: return pd.DataFrame()
    data = gtgt_decl['data']

    detailed_data = {
        'Mã chỉ tiêu': [
            '[21]', '[22]', '[23]', '[24]', '[25]', '[26]', '[29]', '[30]', '[31]',
            '[32]', '[33]', '[32a]', '[34]', '[35]', '[36]', '[37]', '[38]', '[39a]',
            '[40a]', '[40b]', '[40]', '[41]', '[42]', '[43]'
        ],
        'Chỉ tiêu': [
            'Thuế GTGT còn được khấu trừ kỳ trước chưa hết',
            'Thuế GTGT còn được khấu trừ kỳ trước chuyển sang',
            'Giá trị của hàng hóa, dịch vụ mua vào',
            'Thuế GTGT của HHDV mua vào',
            'Thuế GTGT của HHDV mua vào được khấu trừ kỳ này',
            'HHDV bán ra không chịu thuế GTGT',
            'HHDV bán ra chịu thuế suất 0%',
            'Doanh thu HHDV bán ra chịu thuế suất 5%',
            'Thuế GTGT HHDV bán ra chịu thuế suất 5%',
            'Doanh thu HHDV bán ra chịu thuế suất 10%',
            'Thuế GTGT HHDV bán ra chịu thuế suất 10%',
            'HHDV bán ra không phải kê khai, nộp thuế GTGT',
            'Tổng doanh thu HHDV bán ra',
            'Tổng thuế GTGT của HHDV bán ra',
            'Thuế GTGT phát sinh trong kỳ',
            'Điều chỉnh giảm thuế GTGT phải nộp',
            'Điều chỉnh tăng thuế GTGT phải nộp',
            'Thuế GTGT của dự án đầu tư được bù trừ',
            'Thuế GTGT phải nộp của HĐKD',
            'Thuế GTGT mua vào của dự án đầu tư cùng tỉnh',
            'Thuế GTGT còn phải nộp trong kỳ',
            'Thuế GTGT chưa khấu trừ hết kỳ này',
            'Thuế GTGT đề nghị hoàn',
            'Thuế GTGT còn được khấu trừ chuyển kỳ sau',
        ],
        'Số tiền': [
            get_single_value(data, 'ct21'), get_single_value(data, 'ct22'), get_single_value(data, 'ct23'),
            get_single_value(data, 'ct24'), get_single_value(data, 'ct25'), get_single_value(data, 'ct26'),
            get_single_value(data, 'ct29'), get_single_value(data, 'ct30'), get_single_value(data, 'ct31'),
            get_single_value(data, 'ct32'), get_single_value(data, 'ct33'), get_single_value(data, 'ct32a'),
            get_single_value(data, 'ct34'), get_single_value(data, 'ct35'), get_single_value(data, 'ct36'),
            get_single_value(data, 'ct37'), get_single_value(data, 'ct38'), get_single_value(data, 'ct39a'),
            get_single_value(data, 'ct40a'), get_single_value(data, 'ct40b'), get_single_value(data, 'ct40'),
            get_single_value(data, 'ct41'), get_single_value(data, 'ct42'), get_single_value(data, 'ct43'),
        ]
    }
    return pd.DataFrame(detailed_data)


def generate_tndn_main_form_df(data):
    if not data: return pd.DataFrame()

    tndn_main_data = {
        'Mã chỉ tiêu': ['A1', 'B1', 'B2', 'B3', 'B4', 'B7', 'B8', 'B9', 'B10', 'B11', 'B12', 'B13', 'B14', 'C1', 'C2',
                        'C3', 'C4', 'C6', 'C7', 'C8', 'C9', 'C10', 'C11', 'C12', 'C13', 'C14', 'C15', 'C16'],
        'Chỉ tiêu': [
            'A1 - Tổng lợi nhuận kế toán trước thuế TNDN',
            'B1 - Các khoản điều chỉnh tăng tổng lợi nhuận trước thuế',
            'B2 - Các khoản chi không được trừ',
            'B3 - Thuế TNDN đã nộp cho phần thu nhập nhận được ở nước ngoài',
            'B4 - Điều chỉnh tăng doanh thu',
            'B7 - Các khoản điều chỉnh làm tăng lợi nhuận trước thuế khác',
            'B8 - Các khoản điều chỉnh giảm tổng lợi nhuận trước thuế',
            'B9 - Giảm trừ các khoản doanh thu đã điều chỉnh tăng',
            'B10 - Chi phí của phần doanh thu điều chỉnh giảm',
            'B11 - Các khoản điều chỉnh làm giảm lợi nhuận trước thuế khác',
            'B12 - Lợi nhuận từ hoạt động BĐS',
            'B13 - Tổng Thu nhập chịu thuế (TNCT)',
            'B14 - TNCT từ hoạt động sản xuất, kinh doanh',
            'C1 - Thu nhập chịu thuế',
            'C2 - Thu nhập chịu thuế từ HĐSXKD',
            'C3 - Thu nhập được miễn thuế',
            'C4 - Chuyển lỗ và bù trừ lãi, lỗ',
            'C6 - Tổng thu nhập tính thuế (TNTT)',
            'C7 - TNTT từ HĐSXKD',
            'C8 - Thuế TNDN từ HĐSXKD theo thuế suất 20%',
            'C9 - Thuế TNDN phải nộp từ HĐSXKD',
            'C10 - Thuế TNDN của hoạt động BĐS phải nộp',
            'C11 - Thuế TNDN đã nộp ở nước ngoài được trừ trong kỳ tính thuế',
            'C12 - Thuế TNDN đã tạm nộp',
            'C13 - Chênh lệch giữa số thuế TNDN phải nộp và đã tạm nộp',
            'C14 - Thuế TNDN còn phải nộp',
            'C15 - Thuế TNDN nộp thừa',
            'C16 - Tổng số thuế TNDN bù trừ cho các nghĩa vụ khác',
        ],
        'Số tiền': [
            get_single_value(data, 'ctA1'), get_single_value(data, 'ctB1'), get_single_value(data, 'ctB2'),
            get_single_value(data, 'ctB3'), get_single_value(data, 'ctB4'), get_single_value(data, 'ctB7'),
            get_single_value(data, 'ctB8'), get_single_value(data, 'ctB9'), get_single_value(data, 'ctB10'),
            get_single_value(data, 'ctB11'), get_single_value(data, 'ctB12'), get_single_value(data, 'ctB13'),
            get_single_value(data, 'ctB14'), get_single_value(data, 'ctC1'), get_single_value(data, 'ctC2'),
            get_single_value(data, 'ctC3'), get_single_value(data, 'ctC4'), get_single_value(data, 'ctC6'),
            get_single_value(data, 'ctC7'), get_single_value(data, 'ctC8'), get_single_value(data, 'ctC9'),
            get_single_value(data, 'ctC10'), get_single_value(data, 'ctC11'), get_single_value(data, 'ctC12'),
            get_single_value(data, 'ctC13'), get_single_value(data, 'ctC14'), get_single_value(data, 'ctC15'),
            get_single_value(data, 'ctC16'),
        ]
    }
    return pd.DataFrame(tndn_main_data)


def generate_tndn_appendix_03_1a_df(data):
    if not data: return pd.DataFrame()

    appendix_data = {
        'Mã chỉ tiêu': ['[04]', '[05]', '[06]', '[08]', '[09]', '[11]', '[12]', '[13]', '[14]', '[15]', '[16]', '[17]',
                        '[18]', '[19]', '[20]', '[21]', '[22]'],
        'Chỉ tiêu': [
            'Tổng doanh thu bán hàng hóa, dịch vụ',
            'Doanh thu bán hàng hóa, dịch vụ xuất khẩu',
            'Các khoản giảm trừ doanh thu',
            'Doanh thu hoạt động tài chính',
            'Chi phí tài chính',
            'Chi phí sản xuất, kinh doanh hàng hóa, dịch vụ',
            'Giá vốn hàng bán',
            'Chi phí bán hàng',
            'Chi phí quản lý doanh nghiệp',
            'Lợi nhuận thuần từ hoạt động kinh doanh',
            'Thu nhập khác',
            'Chi phí khác',
            'Lợi nhuận khác',
            'Lợi nhuận từ HĐSXKD',
            'Lợi nhuận từ hoạt động chuyển nhượng BĐS',
            'Tổng lợi nhuận kế toán trước thuế TNDN',
            'Trích lập quỹ KH&CN (nếu có)'
        ],
        'Số tiền': [
            get_single_value(data, 'ct04'), get_single_value(data, 'ct05'), get_single_value(data, 'ct06'),
            get_single_value(data, 'ct08'), get_single_value(data, 'ct09'), get_single_value(data, 'ct11'),
            get_single_value(data, 'ct12'), get_single_value(data, 'ct13'), get_single_value(data, 'ct14'),
            get_single_value(data, 'ct15'), get_single_value(data, 'ct16'), get_single_value(data, 'ct17'),
            get_single_value(data, 'ct18'), get_single_value(data, 'ct19'), get_single_value(data, 'ct20'),
            get_single_value(data, 'ct21'), get_single_value(data, 'ct22'),
        ]
    }
    return pd.DataFrame(appendix_data)


def generate_balance_sheet_df(declarations):
    bctc_decl = next((d for d in declarations if 'BCTC' in d['loai_tk']), None)
    if not bctc_decl: return pd.DataFrame()
    data = bctc_decl['data']

    balance_sheet_data = {
        'Chỉ tiêu': [
            'A - TÀI SẢN NGẮN HẠN', 'I. Tiền và các khoản tương đương tiền', 'II. Đầu tư tài chính ngắn hạn',
            'III. Các khoản phải thu ngắn hạn',
            'IV. Hàng tồn kho', 'V. Tài sản ngắn hạn khác', 'B - TÀI SẢN DÀI HẠN', 'I. Các khoản phải thu dài hạn',
            'II. Tài sản cố định', 'III. Bất động sản đầu tư', 'IV. Tài sản dở dang dài hạn',
            'V. Đầu tư tài chính dài hạn',
            'VI. Tài sản dài hạn khác', 'TỔNG CỘNG TÀI SẢN', 'C - NỢ PHẢI TRẢ', 'I. Nợ ngắn hạn', 'II. Nợ dài hạn',
            'D - VỐN CHỦ SỞ HỮU', 'I. Vốn chủ sở hữu', 'TỔNG CỘNG NGUỒN VỐN'
        ],
        'Mã số': [
            '100', '110', '120', '130', '140', '150', '200', '210', '220', '230', '240', '250', '260', '270',
            '300', '310', '330', '400', '410', '440'
        ],
        'Số cuối năm': [
            get_single_value(data, 'scn_ct100'), get_single_value(data, 'scn_ct110'),
            get_single_value(data, 'scn_ct120'), get_single_value(data, 'scn_ct130'),
            get_single_value(data, 'scn_ct140'), get_single_value(data, 'scn_ct150'),
            get_single_value(data, 'scn_ct200'), get_single_value(data, 'scn_ct210'),
            get_single_value(data, 'scn_ct220'), get_single_value(data, 'scn_ct230'),
            get_single_value(data, 'scn_ct240'), get_single_value(data, 'scn_ct250'),
            get_single_value(data, 'scn_ct260'), get_single_value(data, 'scn_ct270'),
            get_single_value(data, 'scn_ct300'), get_single_value(data, 'scn_ct310'),
            get_single_value(data, 'scn_ct330'), get_single_value(data, 'scn_ct400'),
            get_single_value(data, 'scn_ct410'), get_single_value(data, 'scn_ct440')
        ],
        'Số đầu năm': [
            get_single_value(data, 'sdn_ct100'), get_single_value(data, 'sdn_ct110'),
            get_single_value(data, 'sdn_ct120'), get_single_value(data, 'sdn_ct130'),
            get_single_value(data, 'sdn_ct140'), get_single_value(data, 'sdn_ct150'),
            get_single_value(data, 'sdn_ct200'), get_single_value(data, 'sdn_ct210'),
            get_single_value(data, 'sdn_ct220'), get_single_value(data, 'sdn_ct230'),
            get_single_value(data, 'sdn_ct240'), get_single_value(data, 'sdn_ct250'),
            get_single_value(data, 'sdn_ct260'), get_single_value(data, 'sdn_ct270'),
            get_single_value(data, 'sdn_ct300'), get_single_value(data, 'sdn_ct310'),
            get_single_value(data, 'sdn_ct330'), get_single_value(data, 'sdn_ct400'),
            get_single_value(data, 'sdn_ct410'), get_single_value(data, 'sdn_ct440')
        ]
    }
    return pd.DataFrame(balance_sheet_data)


def generate_income_statement_df(declarations):
    bctc_decl = next((d for d in declarations if 'BCTC' in d['loai_tk']), None)
    if not bctc_decl: return pd.DataFrame()
    data = bctc_decl['data']

    income_statement_data = {
        'Chỉ tiêu': [
            '1. Doanh thu bán hàng và cung cấp dịch vụ', '2. Các khoản giảm trừ doanh thu',
            '3. Doanh thu thuần về bán hàng và cung cấp dịch vụ',
            '4. Giá vốn hàng bán', '5. Lợi nhuận gộp về bán hàng và cung cấp dịch vụ',
            '6. Doanh thu hoạt động tài chính',
            '7. Chi phí tài chính', 'Trong đó: Chi phí lãi vay', '8. Chi phí bán hàng',
            '9. Chi phí quản lý doanh nghiệp',
            '10. Lợi nhuận thuần từ hoạt động kinh doanh', '11. Thu nhập khác', '12. Chi phí khác',
            '13. Lợi nhuận khác', '14. Tổng lợi nhuận kế toán trước thuế', '15. Chi phí thuế TNDN hiện hành',
            '16. Chi phí thuế TNDN hoãn lại', '17. Lợi nhuận sau thuế thu nhập doanh nghiệp'
        ],
        'Mã số': [
            '01', '02', '10', '11', '20', '21', '22', '23', '25', '26', '30', '31', '32', '40', '50', '51', '52', '60'
        ],
        'Năm nay': [
            get_single_value(data, 'kqkd_nn_ct01'), get_single_value(data, 'kqkd_nn_ct02'),
            get_single_value(data, 'kqkd_nn_ct10'),
            get_single_value(data, 'kqkd_nn_ct11'), get_single_value(data, 'kqkd_nn_ct20'),
            get_single_value(data, 'kqkd_nn_ct21'),
            get_single_value(data, 'kqkd_nn_ct22'), get_single_value(data, 'kqkd_nn_ct23'),
            get_single_value(data, 'kqkd_nn_ct25'),
            get_single_value(data, 'kqkd_nn_ct26'), get_single_value(data, 'kqkd_nn_ct30'),
            get_single_value(data, 'kqkd_nn_ct31'),
            get_single_value(data, 'kqkd_nn_ct32'), get_single_value(data, 'kqkd_nn_ct40'),
            get_single_value(data, 'kqkd_nn_ct50'),
            get_single_value(data, 'kqkd_nn_ct51'), get_single_value(data, 'kqkd_nn_ct52'),
            get_single_value(data, 'kqkd_nn_ct60')
        ],
        'Năm trước': [
            get_single_value(data, 'kqkd_nt_ct01'), get_single_value(data, 'kqkd_nt_ct02'),
            get_single_value(data, 'kqkd_nt_ct10'),
            get_single_value(data, 'kqkd_nt_ct11'), get_single_value(data, 'kqkd_nt_ct20'),
            get_single_value(data, 'kqkd_nt_ct21'),
            get_single_value(data, 'kqkd_nt_ct22'), get_single_value(data, 'kqkd_nt_ct23'),
            get_single_value(data, 'kqkd_nt_ct25'),
            get_single_value(data, 'kqkd_nt_ct26'), get_single_value(data, 'kqkd_nt_ct30'),
            get_single_value(data, 'kqkd_nt_ct31'),
            get_single_value(data, 'kqkd_nt_ct32'), get_single_value(data, 'kqkd_nt_ct40'),
            get_single_value(data, 'kqkd_nt_ct50'),
            get_single_value(data, 'kqkd_nt_ct51'), get_single_value(data, 'kqkd_nt_ct52'),
            get_single_value(data, 'kqkd_nt_ct60')
        ]
    }
    return pd.DataFrame(income_statement_data)


def generate_trial_balance_df(declarations):
    bctc_decl = next((d for d in declarations if 'BCTC' in d['loai_tk']), None)
    if not bctc_decl: return pd.DataFrame()
    data = bctc_decl['data']

    accounts = {
        "ct111": "Tiền mặt", "ct112": "Tiền gửi ngân hàng", "ct121": "Chứng khoán kinh doanh",
        "ct128": "Đầu tư nắm giữ đến ngày đáo hạn",
        "ct131": "Phải thu của khách hàng", "ct133": "Thuế GTGT được khấu trừ", "ct141": "Tạm ứng",
        "ct152": "Nguyên liệu, vật liệu",
        "ct153": "Công cụ, dụng cụ", "ct154": "Chi phí SX, KD dở dang", "ct155": "Thành phẩm", "ct156": "Hàng hóa",
        "ct157": "Hàng gửi đi bán", "ct211": "TSCĐ hữu hình", "ct214": "Hao mòn TSCĐ", "ct242": "Chi phí trả trước",
        "ct331": "Phải trả cho người bán", "ct333": "Thuế và các khoản phải nộp NN", "ct334": "Phải trả người lao động",
        "ct338": "Phải trả, phải nộp khác", "ct341": "Vay và nợ thuê tài chính", "ct411": "Vốn đầu tư của chủ sở hữu",
        "ct421": "Lợi nhuận sau thuế chưa phân phối", "ct511": "Doanh thu bán hàng và CCDV",
        "ct632": "Giá vốn hàng bán",
        "ct641": "Chi phí bán hàng", "ct642": "Chi phí quản lý doanh nghiệp", "ct711": "Thu nhập khác",
        "ct811": "Chi phí khác",
        "ct911": "Xác định kết quả kinh doanh"
    }

    trial_balance_data = []
    for code, name in accounts.items():
        sddk_no = get_single_value(data, f'sddk_no_{code}', 0)
        sddk_co = get_single_value(data, f'sddk_co_{code}', 0)
        ps_no = get_single_value(data, f'ps_no_{code}', 0)
        ps_co = get_single_value(data, f'ps_co_{code}', 0)
        sdck_no = get_single_value(data, f'sdck_no_{code}', 0)
        sdck_co = get_single_value(data, f'sdck_co_{code}', 0)

        if any([sddk_no, sddk_co, ps_no, ps_co, sdck_no, sdck_co]):
            trial_balance_data.append({
                'Số hiệu TK': code.replace('ct', ''),
                'Tên tài khoản': name,
                'Số dư đầu kỳ - Nợ': sddk_no,
                'Số dư đầu kỳ - Có': sddk_co,
                'Số phát sinh trong kỳ - Nợ': ps_no,
                'Số phát sinh trong kỳ - Có': ps_co,
                'Số dư cuối kỳ - Nợ': sdck_no,
                'Số dư cuối kỳ - Có': sdck_co,
            })

    return pd.DataFrame(trial_balance_data)


def generate_tndn_summary(declarations):
    tndn_decl = next((d for d in declarations if 'TNDN' in d['loai_tk']), None)
    if not tndn_decl: return pd.DataFrame()
    data = tndn_decl['data']
    summary_data = {
        'Chỉ tiêu': ['Tổng DT bán hàng và cung cấp dịch vụ (PL 03-1A)', 'Tổng chi phí (PL 03-1A)',
                     'Lợi nhuận kế toán trước thuế (TK chính)', 'Thu nhập khác (PL 03-1A)',
                     'Thuế TNDN phải nộp (TK chính)'],
        'Số tiền (VND)': [
            get_single_value(data, XML_TAG_MAP['03/TNDN']['tong_doanh_thu_pl'], 0),
            get_single_value(data, XML_TAG_MAP['03/TNDN']['chi_phi'], 0),
            get_single_value(data, XML_TAG_MAP['03/TNDN']['loi_nhuan_truoc_thue'], 0),
            get_single_value(data, XML_TAG_MAP['03/TNDN']['thu_nhap_khac'], 0),
            get_single_value(data, XML_TAG_MAP['03/TNDN']['thue_tndn_phai_nop'], 0),
        ]}
    return pd.DataFrame(summary_data)


def generate_tncn_summary(declarations):
    tncn_decl = next((d for d in declarations if '05/QTT-TNCN' in d['loai_tk']), None)
    if not tncn_decl: return pd.DataFrame(), pd.DataFrame()
    data = tncn_decl['data']
    summary_data = {
        'Chỉ tiêu': ['Tổng số lao động', 'Tổng thu nhập chịu thuế trả cho cá nhân', 'Tổng số thuế TNCN đã khấu trừ'],
        'Số liệu': [
            int(get_single_value(data, XML_TAG_MAP['05/QTT-TNCN']['tong_so_ld'], 0)),
            get_single_value(data, XML_TAG_MAP['05/QTT-TNCN']['tong_tnct'], 0),
            get_single_value(data, XML_TAG_MAP['05/QTT-TNCN']['tong_thue_da_khau_tru'], 0),
        ]}
    summary_df = pd.DataFrame(summary_data)

    def to_list(value):
        if value is None: return []
        return value if isinstance(value, list) else [value]

    details_data = {
        'Họ và tên': to_list(data.get(XML_TAG_MAP['05/QTT-TNCN']['pl_ho_ten'])),
        'Mã số thuế': to_list(data.get(XML_TAG_MAP['05/QTT-TNCN']['pl_mst'])),
        'Tổng TNCT (VND)': [float(x or 0) for x in to_list(data.get(XML_TAG_MAP['05/QTT-TNCN']['pl_tnct']))],
        'Số thuế đã khấu trừ (VND)': [float(x or 0) for x in
                                      to_list(data.get(XML_TAG_MAP['05/QTT-TNCN']['pl_thue_da_khau_tru']))],
    }
    try:
        max_len = max(len(v) for v in details_data.values()) if details_data else 0
        for k, v in details_data.items():
            if len(v) < max_len:
                v.extend([None] * (max_len - len(v)))
        details_df = pd.DataFrame(details_data)
    except (ValueError, TypeError):
        details_df = pd.DataFrame()
    return summary_df, details_df


def generate_tncn_kk_summary(declarations):
    tncn_decls = [d for d in declarations if '05/KK-TNCN' in d['loai_tk']]
    if not tncn_decls: return pd.DataFrame()
    periods = sorted(list(set(d['ky'] for d in tncn_decls)))
    summary_data = {
        'Chỉ tiêu': ['Tổng số người lao động', 'Tổng TNCT trả cho cá nhân', 'Tổng số thuế TNCN đã khấu trừ']}
    for period in periods:
        p_data = next((d['data'] for d in tncn_decls if d['ky'] == period), {})
        tong_ld = get_single_value(p_data, XML_TAG_MAP['05/KK-TNCN']['tong_so_ld'], 0)
        tong_tnct = get_single_value(p_data, XML_TAG_MAP['05/KK-TNCN']['tong_tnct'], 0)
        tong_thue = get_single_value(p_data, XML_TAG_MAP['05/KK-TNCN']['tong_thue_da_khau_tru'], 0)
        summary_data[period] = [tong_ld, tong_tnct, tong_thue]
    df = pd.DataFrame(summary_data)
    if len(periods) > 0: df['Tổng cộng'] = df[periods].sum(axis=1)
    return df


def run_risk_checks(declarations, gtgt_summary_df, tncn_kk_summary_df, accounting_standard,
                    output_invoice_pre_tax_total=None, input_invoice_data=None):
    results = []
    tndn_decl = next((d for d in declarations if 'TNDN' in d['loai_tk']), None)
    bctc_decl = next((d for d in declarations if 'BCTC' in d['loai_tk']), None)
    tncn_qtt_decl = next((d for d in declarations if '05/QTT-TNCN' in d['loai_tk']), None)

    tong_dt_gtgt = 0
    dt_tndn = 0
    if tndn_decl and not gtgt_summary_df.empty:
        tong_dt_gtgt_series = gtgt_summary_df.loc[
            gtgt_summary_df['Chỉ tiêu'] == 'Doanh thu chịu thuế GTGT', 'Tổng cộng']
        tong_dt_gtgt = tong_dt_gtgt_series.iloc[0] if not tong_dt_gtgt_series.empty else 0
        dt_tndn_pl = get_single_value(tndn_decl['data'], XML_TAG_MAP['03/TNDN']['tong_doanh_thu_pl'], 0)
        tn_khac_tndn = get_single_value(tndn_decl['data'], XML_TAG_MAP['03/TNDN']['thu_nhap_khac'], 0)
        dt_tndn = dt_tndn_pl + tn_khac_tndn
        chenh_lech = tong_dt_gtgt - dt_tndn
        results.append({
            "Nội dung": "Doanh thu GTGT vs. Doanh thu TNDN",
            "Số liệu A": f"{tong_dt_gtgt:,.0f} (TK GTGT)",
            "Số liệu B": f"{dt_tndn:,.0f} (QT TNDN)",
            "Chênh lệch": f"{chenh_lech:,.0f}",
            "Trạng thái": "Cảnh báo" if chenh_lech != 0 else "Khớp",
            "Gợi ý": "Đối chiếu DT bán ra trên tờ khai GTGT và QT TNDN."
        })
    else:
        results.append({"Nội dung": "Doanh thu GTGT vs. Doanh thu TNDN", "Số liệu A": "N/A", "Số liệu B": "N/A",
                        "Chênh lệch": "N/A", "Trạng thái": "Không đủ dữ liệu",
                        "Gợi ý": "Cần tải lên cả TK GTGT năm và QT TNDN."})

    if output_invoice_pre_tax_total is not None:
        if not gtgt_summary_df.empty and tong_dt_gtgt > 0:
            chenh_lech_gtgt = tong_dt_gtgt - output_invoice_pre_tax_total
            results.append({
                "Nội dung": "Doanh thu GTGT vs. Bảng kê hóa đơn",
                "Số liệu A": f"{tong_dt_gtgt:,.0f} (TK GTGT)",
                "Số liệu B": f"{output_invoice_pre_tax_total:,.0f} (Bảng kê HĐ)",
                "Chênh lệch": f"{chenh_lech_gtgt:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_gtgt != 0 else "Khớp",
                "Gợi ý": "Kiểm tra chênh lệch giữa tổng doanh thu trên các tờ khai GTGT và tổng doanh thu từ bảng kê hóa đơn bán ra."
            })
        if tndn_decl and dt_tndn > 0:
            chenh_lech_tndn = dt_tndn - output_invoice_pre_tax_total
            results.append({
                "Nội dung": "Doanh thu TNDN vs. Bảng kê hóa đơn",
                "Số liệu A": f"{dt_tndn:,.0f} (QT TNDN)",
                "Số liệu B": f"{output_invoice_pre_tax_total:,.0f} (Bảng kê HĐ)",
                "Chênh lệch": f"{chenh_lech_tndn:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_tndn != 0 else "Khớp",
                "Gợi ý": "Kiểm tra chênh lệch giữa tổng doanh thu trên Quyết toán TNDN và tổng doanh thu từ bảng kê hóa đơn bán ra."
            })
    else:
        results.append({"Nội dung": "Đối chiếu Bảng kê hóa đơn bán ra", "Số liệu A": "N/A", "Số liệu B": "N/A",
                        "Chênh lệch": "N/A", "Trạng thái": "Không đủ dữ liệu",
                        "Gợi ý": "Hãy tải lên Bảng kê hóa đơn bán ra để thực hiện đối chiếu."})

    if tncn_qtt_decl and not tncn_kk_summary_df.empty:
        tong_thue_kk_series = tncn_kk_summary_df.loc[
            tncn_kk_summary_df['Chỉ tiêu'] == 'Tổng số thuế TNCN đã khấu trừ', 'Tổng cộng']
        tong_thue_kk = tong_thue_kk_series.iloc[0] if not tong_thue_kk_series.empty else 0

        thue_qtt = get_single_value(tncn_qtt_decl['data'], XML_TAG_MAP['05/QTT-TNCN']['tong_thue_da_khau_tru'], 0)
        chenh_lech_tncn = tong_thue_kk - thue_qtt
        results.append({
            "Nội dung": "Đối chiếu thuế TNCN khấu trừ (Khai kỳ vs. Quyết toán)",
            "Số liệu A": f"{tong_thue_kk:,.0f} (Tổng các kỳ)",
            "Số liệu B": f"{thue_qtt:,.0f} (Quyết toán năm)",
            "Chênh lệch": f"{chenh_lech_tncn:,.0f}",
            "Trạng thái": "Cảnh báo" if chenh_lech_tncn != 0 else "Khớp",
            "Gợi ý": "Đối chiếu tổng số thuế TNCN đã khấu trừ trên các tờ khai 05/KK-TNCN với chỉ tiêu [31] trên tờ khai 05/QTT-TNCN."
        })
    else:
        results.append({"Nội dung": "Đối chiếu thuế TNCN khấu trừ (Khai kỳ vs. Quyết toán)", "Số liệu A": "N/A",
                        "Số liệu B": "N/A",
                        "Chênh lệch": "N/A", "Trạng thái": "Không đủ dữ liệu",
                        "Gợi ý": "Cần tải lên cả tờ khai 05/KK-TNCN (tháng/quý) và tờ khai 05/QTT-TNCN (năm)."})

    # CẬP NHẬT: Logic đối chiếu thuế GTGT đầu vào
    gtgt_decls_data = [d['data'] for d in declarations if '01/GTGT' in d['loai_tk']]
    if input_invoice_data and gtgt_decls_data:
        total_ct23_from_tk = sum(get_single_value(data, 'ct23') for data in gtgt_decls_data)
        total_ct24_from_tk = sum(get_single_value(data, 'ct24') for data in gtgt_decls_data)
        total_ct25_from_tk = sum(get_single_value(data, 'ct25') for data in gtgt_decls_data)

        total_pre_tax_from_invoice = input_invoice_data.get('valid_summary', {}).get('total_pre_tax', 0)
        total_tax_from_invoice = input_invoice_data.get('valid_summary', {}).get('total_tax', 0)

        chenh_lech_ct23 = total_ct23_from_tk - total_pre_tax_from_invoice
        results.append({
            "Nội dung": "GTGT đầu vào: HHDV mua vào (TK vs Bảng kê)",
            "Số liệu A": f"{total_ct23_from_tk:,.0f} (TK GTGT - CT23)",
            "Số liệu B": f"{total_pre_tax_from_invoice:,.0f} (Bảng kê HĐ vào)",
            "Chênh lệch": f"{chenh_lech_ct23:,.0f}",
            "Trạng thái": "Cảnh báo" if abs(chenh_lech_ct23) > 1 else "Khớp",
            "Gợi ý": "Đối chiếu tổng giá trị HHDV mua vào trên các tờ khai GTGT với tổng tiền chưa thuế trên bảng kê hóa đơn đầu vào."
        })

        chenh_lech_24_25 = total_ct24_from_tk - total_ct25_from_tk
        results.append({
            "Nội dung": "GTGT đầu vào: Thuế mua vào vs Thuế được khấu trừ (trên TK)",
            "Số liệu A": f"{total_ct24_from_tk:,.0f} (TK GTGT - CT24)",
            "Số liệu B": f"{total_ct25_from_tk:,.0f} (TK GTGT - CT25)",
            "Chênh lệch": f"{chenh_lech_24_25:,.0f}",
            "Trạng thái": "Cảnh báo" if chenh_lech_24_25 != 0 else "OK",
            "Gợi ý": "Kiểm tra lý do không được khấu trừ toàn bộ thuế GTGT đầu vào (nếu có)."
        })

        chenh_lech_ct25 = total_ct25_from_tk - total_tax_from_invoice
        results.append({
            "Nội dung": "GTGT đầu vào: Thuế được khấu trừ (TK vs Bảng kê)",
            "Số liệu A": f"{total_ct25_from_tk:,.0f} (TK GTGT - CT25)",
            "Số liệu B": f"{total_tax_from_invoice:,.0f} (Bảng kê HĐ vào)",
            "Chênh lệch": f"{chenh_lech_ct25:,.0f}",
            "Trạng thái": "Cảnh báo" if abs(chenh_lech_ct25) > 1 else "Khớp",
            "Gợi ý": "Đối chiếu tổng thuế GTGT được khấu trừ trên các tờ khai với tổng tiền thuế trên bảng kê hóa đơn đầu vào."
        })
    else:
        results.append({"Nội dung": "Đối chiếu thuế GTGT đầu vào", "Số liệu A": "N/A", "Số liệu B": "N/A",
                        "Chênh lệch": "N/A", "Trạng thái": "Không đủ dữ liệu",
                        "Gợi ý": "Cần tải lên cả tờ khai GTGT và Bảng kê hóa đơn đầu vào."})

    if bctc_decl:
        if accounting_standard == "Chưa chọn":
            results.append({"Nội dung": "Phân tích BCTC", "Số liệu A": "N/A", "Số liệu B": "N/A", "Chênh lệch": "N/A",
                            "Trạng thái": "Không đủ dữ liệu",
                            "Gợi ý": "Vui lòng chọn Chế độ kế toán (TT133/TT200) để thực hiện đối chiếu BCTC."})

        elif accounting_standard == "Thông tư 133":
            bctc_data = bctc_decl.get('data', {})
            doanh_thu_kqkd = get_single_value(bctc_data, 'kqkd_nn_ct01', 0)
            tien_mat = get_single_value(bctc_data, 'scn_ct110', 0)
            chi_phi_lai_vay = get_single_value(bctc_data, 'kqkd_nn_ct23', 0)
            is_risk = tien_mat > 1000000000 and chi_phi_lai_vay > 0
            results.append({
                "Nội dung": "Chi phí lãi vay bất thường",
                "Số liệu A": f"{tien_mat:,.0f} (Tiền)",
                "Số liệu B": f"{chi_phi_lai_vay:,.0f} (CP Lãi vay)",
                "Chênh lệch": "N/A",
                "Trạng thái": "Cảnh báo" if is_risk else "OK",
                "Gợi ý": "Xem xét tính hợp lý khi có lượng tiền mặt lớn nhưng vẫn đi vay."
            })

            ps_no_131 = get_single_value(bctc_data, 'ps_no_ct131', 0)
            ps_co_511 = get_single_value(bctc_data, 'ps_co_ct511', 0)
            ps_co_3331 = get_single_value(bctc_data, 'ps_co_ct3331', 0)
            ps_co_711 = get_single_value(bctc_data, 'ps_co_ct711', 0)
            tong_co = ps_co_511 + ps_co_3331 + ps_co_711
            chenh_lech_131 = ps_no_131 - tong_co
            goi_y_131 = "OK"
            if chenh_lech_131 > 0:
                goi_y_131 = "PS Nợ 131 > PS Có (511+3331+711). Rủi ro ghi nhận thiếu doanh thu."
            elif chenh_lech_131 < 0:
                goi_y_131 = "PS Nợ 131 < PS Có (511+3331+711). Yêu cầu làm rõ, đối chiếu TK đối ứng."
            results.append({
                "Nội dung": "Đối chiếu PS Nợ TK 131",
                "Số liệu A": f"{ps_no_131:,.0f} (PS Nợ 131)",
                "Số liệu B": f"{tong_co:,.0f} (PS Có 511+3331+711)",
                "Chênh lệch": f"{chenh_lech_131:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_131 != 0 else "Khớp",
                "Gợi ý": goi_y_131
            })

            ps_co_512 = get_single_value(bctc_data, 'ps_co_ct512', 0)
            tong_ps_co_dt = ps_co_511 + ps_co_512
            chenh_lech_dt = doanh_thu_kqkd - tong_ps_co_dt
            results.append({
                "Nội dung": "Doanh thu trên KQKD vs CĐTK",
                "Số liệu A": f"{doanh_thu_kqkd:,.0f} (KQKD)",
                "Số liệu B": f"{tong_ps_co_dt:,.0f} (PS Có 511+512)",
                "Chênh lệch": f"{chenh_lech_dt:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_dt != 0 else "Khớp",
                "Gợi ý": "Đối chiếu số liệu doanh thu giữa các phụ lục BCTC."
            })

            ps_co_154 = get_single_value(bctc_data, 'ps_co_ct154', 0)
            ps_no_155 = get_single_value(bctc_data, 'ps_no_ct155', 0)
            chenh_lech_154_155 = ps_no_155 - ps_co_154
            goi_y_154_155 = "OK"
            if chenh_lech_154_155 != 0:
                if chenh_lech_154_155 > 0:
                    goi_y_154_155 = "PS Nợ 155 > PS Có 154. Yêu cầu đối chiếu với PS Nợ 632 và PS Có 511/512."
                else:
                    goi_y_154_155 = "PS Có 154 > PS Nợ 155. Dấu hiệu bán hàng không nhập kho, biếu tặng không ghi nhận doanh thu."
            results.append({
                "Nội dung": "Chi phí dở dang vs. Thành phẩm (TK 154 vs 155)",
                "Số liệu A": f"{ps_co_154:,.0f} (PS Có 154)",
                "Số liệu B": f"{ps_no_155:,.0f} (PS Nợ 155)",
                "Chênh lệch": f"{chenh_lech_154_155:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_154_155 != 0 else "Khớp",
                "Gợi ý": goi_y_154_155
            })

            ps_no_621 = get_single_value(bctc_data, 'ps_no_ct621', 0)
            ps_no_622 = get_single_value(bctc_data, 'ps_no_ct622', 0)
            ps_no_627 = get_single_value(bctc_data, 'ps_no_ct627', 0)
            tong_chi_phi_sx = ps_no_621 + ps_no_622 + ps_no_627
            chenh_lech_cpsx = ps_co_154 - tong_chi_phi_sx
            results.append({
                "Nội dung": "Kết chuyển chi phí SX (TK 154 vs 621, 622, 627)",
                "Số liệu A": f"{ps_co_154:,.0f} (PS Có 154)",
                "Số liệu B": f"{tong_chi_phi_sx:,.0f} (PS Nợ 621+622+627)",
                "Chênh lệch": f"{chenh_lech_cpsx:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_cpsx != 0 else "Khớp",
                "Gợi ý": "Kiểm tra việc kết chuyển chi phí sản xuất vào TK 154."
            })

            hang_ban_tra_lai = get_single_value(bctc_data, 'kqkd_nn_ct02', 0)
            results.append({
                "Nội dung": "Hàng bán trả lại",
                "Số liệu A": f"{hang_ban_tra_lai:,.0f} (Giảm trừ DT)",
                "Số liệu B": "N/A", "Chênh lệch": "N/A",
                "Trạng thái": "Cảnh báo" if hang_ban_tra_lai > 0 else "OK",
                "Gợi ý": "Có phát sinh hàng bán trả lại. Yêu cầu cung cấp chi tiết PS Có TK 632 để kiểm tra việc giảm giá vốn." if hang_ban_tra_lai > 0 else "Không có hàng bán trả lại."
            })

            ps_co_152 = get_single_value(bctc_data, 'ps_co_ct152', 0)
            chenh_lech_152_621 = ps_co_152 - ps_no_621
            goi_y_152_621 = "OK"
            if chenh_lech_152_621 < 0:
                goi_y_152_621 = "PS Có 152 < PS Nợ 621. Nghi vấn ghi nhận chi phí NVL nhưng không xuất kho (thiếu hóa đơn)."
            elif chenh_lech_152_621 > 0:
                goi_y_152_621 = "PS Có 152 > PS Nợ 621. Dấu hiệu xuất NVL để trao đổi/bán không ghi nhận doanh thu."
            results.append({
                "Nội dung": "Xuất kho NVL vs. Chi phí NVL (TK 152 vs 621)",
                "Số liệu A": f"{ps_co_152:,.0f} (PS Có 152)",
                "Số liệu B": f"{ps_no_621:,.0f} (PS Nợ 621)",
                "Chênh lệch": f"{chenh_lech_152_621:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_152_621 != 0 else "Khớp",
                "Gợi ý": goi_y_152_621
            })

            htk_dau_ky = get_single_value(bctc_data, 'sdn_ct140', 0)
            htk_cuoi_ky = get_single_value(bctc_data, 'scn_ct140', 0)
            goi_y_htk = "OK"
            trang_thai_htk = "OK"
            if doanh_thu_kqkd > 0 and htk_cuoi_ky > (doanh_thu_kqkd * 2):
                trang_thai_htk = "Cảnh báo"
                goi_y_htk = f"HTK cuối kỳ ({htk_cuoi_ky:,.0f}) gấp {(htk_cuoi_ky / doanh_thu_kqkd):.1f} lần doanh thu. "
            if htk_cuoi_ky >= htk_dau_ky and htk_dau_ky > 0:
                if trang_thai_htk != "Cảnh báo": goi_y_htk = ""
                trang_thai_htk = "Cảnh báo"
                goi_y_htk += "HTK không giảm hoặc tăng so với đầu kỳ. Dấu hiệu tồn kho ảo/kém luân chuyển."
            results.append({
                "Nội dung": "Rủi ro hàng tồn kho ảo",
                "Số liệu A": f"{htk_dau_ky:,.0f} (HTK Đầu kỳ)",
                "Số liệu B": f"{htk_cuoi_ky:,.0f} (HTK Cuối kỳ)",
                "Chênh lệch": f"{htk_cuoi_ky - htk_dau_ky:,.0f}",
                "Trạng thái": trang_thai_htk,
                "Gợi ý": goi_y_htk
            })

            balance_checks = [
                {'tk': 'sdck_co_ct131', 'name': 'Dư Có TK 131',
                 'hint': 'Kiểm tra chi tiết: Nếu là người mua trả trước, kiểm tra hợp đồng. Nếu hàng đã tiêu thụ, phải ghi nhận doanh thu tính thuế.'},
                {'tk': 'sdck_co_ct337', 'name': 'Dư Có TK 337 (Hợp đồng XD)',
                 'hint': 'Kiểm tra hợp đồng, tiến độ để ghi nhận doanh thu theo hạng mục hoàn thành.'},
                {'tk': 'sdck_co_ct3387', 'name': 'Dư Có TK 3387 (DT chưa thực hiện)',
                 'hint': 'Kiểm tra chi tiết: Nếu hàng đã giao, BĐS đã bàn giao, hạng mục XD đã hoàn thành, phải ghi nhận doanh thu tính thuế.'},
                {'tk': 'sdck_no_ct157', 'name': 'Dư Nợ TK 157 (Hàng gửi bán)',
                 'hint': 'Kiểm tra chi tiết: Nếu hàng đã gửi cho khách hàng, phải ghi nhận doanh thu tính thuế theo quy định.'},
                {'tk': 'sdck_no_ct136', 'name': 'Dư Nợ TK 136 (Phải thu nội bộ)',
                 'hint': 'Kiểm tra chi tiết: Nếu là tiền bán hàng nội bộ, phải ghi nhận doanh thu tính thuế.'},
                {'tk': 'sdck_no_ct138', 'name': 'Dư Nợ TK 138 (Phải thu khác)',
                 'hint': 'Kiểm tra chi tiết khoản phải thu khác. Nếu là giao dịch hàng hóa/dịch vụ, phải ghi nhận doanh thu.'},
                {'tk': 'sdck_co_ct138', 'name': 'Dư Có TK 138 (Phải trả khác)',
                 'hint': 'Kiểm tra chi tiết khoản thu thừa. Nếu là giao dịch hàng hóa/dịch vụ, phải ghi nhận doanh thu.'},
            ]
            for check in balance_checks:
                balance = get_single_value(bctc_data, check['tk'], 0)
                results.append({
                    "Nội dung": check['name'],
                    "Số liệu A": f"{balance:,.0f}",
                    "Số liệu B": "N/A", "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo" if balance > 0 else "OK",
                    "Gợi ý": check['hint'] if balance > 0 else "Không có số dư bất thường."
                })

            ps_co_155 = get_single_value(bctc_data, 'ps_co_ct155', 0)
            ps_co_156 = get_single_value(bctc_data, 'ps_co_ct156', 0)
            ps_no_632 = get_single_value(bctc_data, 'ps_no_ct632', 0)
            tong_xuat_kho = ps_co_155 + ps_co_156
            chenh_lech_gvon = tong_xuat_kho - ps_no_632
            goi_y_gvon = "OK"
            if chenh_lech_gvon != 0:
                if chenh_lech_gvon > 0:
                    goi_y_gvon = "Xuất kho > Giá vốn. Dấu hiệu xuất tiêu thụ/biếu tặng không ghi nhận giá vốn, hoặc hàng bán trả lại không giảm giá vốn."
                else:
                    goi_y_gvon = "Xuất kho < Giá vốn. Có thể do dự phòng giảm giá HTK. Nếu không, có thể là hàng mua bán thẳng không qua kho. Cần đối chiếu PS Có TK 511/512."
            results.append({
                "Nội dung": "Xuất kho TP, HH vs. Giá vốn (TK 155, 156 vs 632)",
                "Số liệu A": f"{tong_xuat_kho:,.0f} (PS Có 155+156)",
                "Số liệu B": f"{ps_no_632:,.0f} (PS Nợ 632)",
                "Chênh lệch": f"{chenh_lech_gvon:,.0f}",
                "Trạng thái": "Cảnh báo" if chenh_lech_gvon != 0 else "Khớp",
                "Gợi ý": goi_y_gvon
            })

            ps_no_2293 = get_single_value(bctc_data, 'ps_no_ct2293', 0)
            ps_co_2293 = get_single_value(bctc_data, 'ps_co_ct2293', 0)
            if ps_no_2293 > 0 or ps_co_2293 > 0:
                results.append({
                    "Nội dung": "Dự phòng phải thu khó đòi (TK 2293/139)",
                    "Số liệu A": f"{ps_no_2293:,.0f} (Trích lập)",
                    "Số liệu B": f"{ps_co_2293:,.0f} (Hoàn nhập)",
                    "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh dự phòng phải thu khó đòi. Yêu cầu kiểm tra sự tương ứng với sự tăng/giảm của các khoản phải thu."
                })

            balance_checks_provision = [
                {'tk': 'sdck_co_ct335', 'name': 'Dư Có TK 335 (Chi phí phải trả)',
                 'hint': 'Kiểm tra chi tiết, đặc biệt là dự phòng bảo hành công trình xây lắp đã hết thời hạn mà chưa hoàn nhập.'},
                {'tk': 'sdck_co_ct352', 'name': 'Dư Có TK 352 (Dự phòng phải trả)',
                 'hint': 'Kiểm tra chi tiết các khoản dự phòng đã trích lập nhưng không sử dụng hoặc không dùng hết mà chưa hoàn nhập.'},
                {'tk': 'sdck_no_ct242', 'name': 'Dư Nợ TK 242 (Chi phí trả trước)',
                 'hint': 'Kiểm tra chi tiết các khoản chi phí trả trước để đảm bảo phân bổ đúng kỳ.'},
            ]
            for check in balance_checks_provision:
                balance = get_single_value(bctc_data, check['tk'], 0)
                if balance > 0:
                    results.append({
                        "Nội dung": check['name'],
                        "Số liệu A": f"{balance:,.0f}", "Số liệu B": "N/A", "Chênh lệch": "N/A",
                        "Trạng thái": "Cảnh báo", "Gợi ý": check['hint']
                    })

            ps_no_811 = get_single_value(bctc_data, 'ps_no_ct811', 0)
            ps_co_711 = get_single_value(bctc_data, 'ps_co_ct711', 0)
            if ps_no_811 > 0:
                results.append({
                    "Nội dung": "Rủi ro thanh lý TSCĐ, bán phế liệu",
                    "Số liệu A": f"{ps_co_711:,.0f} (Thu nhập khác)",
                    "Số liệu B": f"{ps_no_811:,.0f} (Chi phí khác)",
                    "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh Chi phí khác. Yêu cầu cung cấp chi tiết TK 711 và 811 để kiểm tra việc hạch toán thu nhập từ thanh lý TSCĐ, bán phế liệu."
                })

        elif accounting_standard == "Thông tư 200":
            bctc_data = bctc_decl.get('data', {})

            tien_mat = get_single_value(bctc_data, 'scn_ct110', 0)
            chi_phi_lai_vay = get_single_value(bctc_data, 'kqkd_nn_ct23', 0)
            is_risk_tien_lai_vay = tien_mat > 1000000000 and chi_phi_lai_vay > 0
            results.append({
                "Nội dung": "Rủi ro Tiền và Chi phí lãi vay (TT200)",
                "Số liệu A": f"{tien_mat:,.0f} (Tiền và TĐT - Mã 110)",
                "Số liệu B": f"{chi_phi_lai_vay:,.0f} (CP lãi vay - Mã 23)",
                "Chênh lệch": "N/A",
                "Trạng thái": "Cảnh báo" if is_risk_tien_lai_vay else "OK",
                "Gợi ý": "Lượng tiền mặt lớn nhưng vẫn phát sinh chi phí lãi vay. Cần xem xét tính hợp lý của các khoản vay."
            })

            phai_thu_noi_bo = get_single_value(bctc_data, 'scn_ct133', 0)
            if phai_thu_noi_bo > 0:
                results.append({
                    "Nội dung": "Rủi ro Phải thu nội bộ ngắn hạn (TT200)",
                    "Số liệu A": f"{phai_thu_noi_bo:,.0f} (Mã 133)",
                    "Số liệu B": "N/A", "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh Phải thu nội bộ. Cần kiểm tra chi tiết các giao dịch trong tập đoàn/công ty mẹ-con để tránh bỏ sót doanh thu."
                })

            phai_thu_khac = get_single_value(bctc_data, 'scn_ct136', 0)
            if phai_thu_khac > 0:
                results.append({
                    "Nội dung": "Rủi ro Phải thu ngắn hạn khác (TT200)",
                    "Số liệu A": f"{phai_thu_khac:,.0f} (Mã 136)",
                    "Số liệu B": "N/A", "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh Phải thu ngắn hạn khác. Cần kiểm tra chi tiết để đảm bảo không có doanh thu bị ghi nhận sai vào khoản mục này."
                })

            provision_checks = [
                {'code': 'ct137', 'name': 'Dự phòng phải thu ngắn hạn khó đòi (*)'},
                {'code': 'ct149', 'name': 'Dự phòng giảm giá hàng tồn kho (*)'},
                {'code': 'ct219', 'name': 'Dự phòng phải thu dài hạn khó đòi (*)'},
                {'code': 'ct321', 'name': 'Quỹ khen thưởng, phúc lợi (*)'},
            ]
            for check in provision_checks:
                balance_sdn = get_single_value(bctc_data, f"sdn_{check['code']}", 0)
                balance_scn = get_single_value(bctc_data, f"scn_{check['code']}", 0)
                if balance_sdn > 0 or balance_scn > 0:
                    results.append({
                        "Nội dung": f"Kiểm tra khoản mục Dự phòng (TT200) - {check['name']}",
                        "Số liệu A": f"{balance_sdn:,.0f} (Đầu năm)",
                        "Số liệu B": f"{balance_scn:,.0f} (Cuối năm)",
                        "Chênh lệch": "N/A",
                        "Trạng thái": "Cảnh báo",
                        "Gợi ý": "Có phát sinh số dư các khoản dự phòng. Yêu cầu DN giải trình về việc trích lập có đúng quy định, hồ sơ kèm theo,..."
                    })

            vat_decls_in_year = [d for d in declarations if '01/GTGT' in d['loai_tk']]
            if vat_decls_in_year:
                last_vat_decl = max(vat_decls_in_year, key=lambda x: x['ky'])
                vat_ct43 = get_single_value(last_vat_decl['data'], 'ct43', 0)
                bctc_ct152 = get_single_value(bctc_data, 'scn_ct152', 0)
                chenh_lech_vat = bctc_ct152 - vat_ct43
                results.append({
                    "Nội dung": "Đối chiếu Thuế GTGT được khấu trừ (TT200)",
                    "Số liệu A": f"{bctc_ct152:,.0f} (BCTC - Mã 152)",
                    "Số liệu B": f"{vat_ct43:,.0f} (TK GTGT cuối kỳ - Chỉ tiêu 43)",
                    "Chênh lệch": f"{chenh_lech_vat:,.0f}",
                    "Trạng thái": "Cảnh báo" if chenh_lech_vat != 0 else "Khớp",
                    "Gợi ý": "Số dư Thuế GTGT được khấu trừ trên BCTC không khớp với số thuế còn được khấu trừ chuyển kỳ sau trên tờ khai GTGT cuối cùng của năm. Yêu cầu DN giải trình."
                })
            else:
                results.append({
                    "Nội dung": "Đối chiếu Thuế GTGT được khấu trừ (TT200)",
                    "Số liệu A": "N/A", "Số liệu B": "N/A", "Chênh lệch": "N/A",
                    "Trạng thái": "Không đủ dữ liệu",
                    "Gợi ý": "Cần tải lên tờ khai GTGT của kỳ cuối cùng trong năm để thực hiện đối chiếu."
                })

            dau_tu_cty_con = get_single_value(bctc_data, 'scn_ct251', 0)
            dau_tu_lkld = get_single_value(bctc_data, 'scn_ct252', 0)
            if dau_tu_cty_con > 0 or dau_tu_lkld > 0:
                results.append({
                    "Nội dung": "Kiểm tra Giao dịch liên kết (TT200)",
                    "Số liệu A": f"{dau_tu_cty_con:,.0f} (Đầu tư vào cty con)",
                    "Số liệu B": f"{dau_tu_lkld:,.0f} (Đầu tư vào LKLĐ)",
                    "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh các khoản đầu tư vào công ty con/liên kết, liên doanh. Cần kiểm tra xem doanh nghiệp có kê khai Phụ lục giao dịch liên kết kèm theo Tờ khai quyết toán thuế TNDN hay không."
                })

            prepayment_short = get_single_value(bctc_data, 'scn_ct312', 0)
            prepayment_long = get_single_value(bctc_data, 'scn_ct332', 0)
            if prepayment_short > 0 or prepayment_long > 0:
                results.append({
                    "Nội dung": "Rủi ro Người mua trả tiền trước (TT200)",
                    "Số liệu A": f"{prepayment_short:,.0f} (Ngắn hạn - Mã 312)",
                    "Số liệu B": f"{prepayment_long:,.0f} (Dài hạn - Mã 332)",
                    "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh khoản người mua trả tiền trước. Yêu cầu DN giải trình chi tiết, có nguy cơ bỏ sót doanh thu."
                })

            payable_employees = get_single_value(bctc_data, 'scn_ct314', 0)
            if payable_employees > 0:
                results.append({
                    "Nội dung": "Rủi ro Phải trả người lao động (TT200)",
                    "Số liệu A": f"{payable_employees:,.0f} (Mã 314)",
                    "Số liệu B": "N/A",
                    "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Yêu cầu DN giải trình đến 31/3 năm sau đã chi hết chưa, cung cấp chứng từ. Nếu chi không hết, có thể bị xuất toán chi phí."
                })

            unearned_revenue_short = get_single_value(bctc_data, 'scn_ct318', 0)
            unearned_revenue_long = get_single_value(bctc_data, 'scn_ct336', 0)
            if unearned_revenue_short > 0 or unearned_revenue_long > 0:
                results.append({
                    "Nội dung": "Rủi ro Doanh thu chưa thực hiện (TT200)",
                    "Số liệu A": f"{unearned_revenue_short:,.0f} (Ngắn hạn - Mã 318)",
                    "Số liệu B": f"{unearned_revenue_long:,.0f} (Dài hạn - Mã 336)",
                    "Chênh lệch": "N/A",
                    "Trạng thái": "Cảnh báo",
                    "Gợi ý": "Có phát sinh doanh thu chưa thực hiện. Yêu cầu DN cung cấp chi tiết để đảm bảo đã ghi nhận đầy đủ doanh thu."
                })

            cplv_kqkd = get_single_value(bctc_data, 'kqkd_nn_ct23', 0)
            tien_chi_tra_lai_vay = abs(get_single_value(bctc_data, 'lctt_nn_ct04', 0))

            trang_thai_lctt = "Không đủ dữ liệu"
            chenh_lech_lctt = "N/A"
            if tien_chi_tra_lai_vay > 0:
                chenh_lech_num = cplv_kqkd - tien_chi_tra_lai_vay
                chenh_lech_lctt = f"{chenh_lech_num:,.0f}"
                trang_thai_lctt = "Cảnh báo" if chenh_lech_num != 0 else "Khớp"

            results.append({
                "Nội dung": "Đối chiếu Chi phí lãi vay (TT200)",
                "Số liệu A": f"{cplv_kqkd:,.0f} (CP lãi vay trên KQKD)",
                "Số liệu B": f"{tien_chi_tra_lai_vay:,.0f} (Tiền trả lãi vay trên LCTT)" if tien_chi_tra_lai_vay > 0 else "Chưa có dữ liệu",
                "Chênh lệch": chenh_lech_lctt,
                "Trạng thái": trang_thai_lctt,
                "Gợi ý": "Đối chiếu chi phí lãi vay trên KQKD và LCTT. Cần tải file XML có chứa dữ liệu LCTT để phân tích."
            })

            items_to_check = {
                "Doanh thu bán hàng": ('kqkd_nn_ct01', 'kqkd_nt_ct01'),
                "Giá vốn hàng bán": ('kqkd_nn_ct11', 'kqkd_nt_ct11'),
                "Chi phí bán hàng": ('kqkd_nn_ct25', 'kqkd_nt_ct25'),
                "Chi phí QLDN": ('kqkd_nn_ct26', 'kqkd_nt_ct26'),
                "Thu nhập khác": ('kqkd_nn_ct31', 'kqkd_nt_ct31'),
            }
            for name, (current_key, prior_key) in items_to_check.items():
                current_val = get_single_value(bctc_data, current_key, 0)
                prior_val = get_single_value(bctc_data, prior_key, 0)

                if prior_val != 0:
                    change_pct = ((current_val - prior_val) / prior_val) * 100
                    if abs(change_pct) > 30:
                        results.append({
                            "Nội dung": f"Biến động bất thường - {name}",
                            "Số liệu A": f"{prior_val:,.0f} (Năm trước)",
                            "Số liệu B": f"{current_val:,.0f} (Năm nay)",
                            "Chênh lệch": f"{change_pct:,.2f}%",
                            "Trạng thái": "Cảnh báo",
                            "Gợi ý": "Yêu cầu Doanh nghiệp giải trình về biến động tăng/giảm đột biến (>30%) so với cùng kỳ."
                        })

    else:
        results.append({"Nội dung": "Phân tích BCTC", "Số liệu A": "N/A", "Số liệu B": "N/A", "Chênh lệch": "N/A",
                        "Trạng thái": "Không đủ dữ liệu", "Gợi ý": "Cần tải lên file Báo cáo tài chính."})

    return results


async def get_gemini_analysis(api_key, dfs_dict, risks_df, notes_content=None):
    prompt = "Bạn là một chuyên gia phân tích thuế. Dựa trên các số liệu tổng hợp từ hồ sơ khai thuế và các tài liệu dưới đây, hãy đưa ra một nhận xét ngắn gọn (khoảng 3-4 gạch đầu dòng) về tình hình tài chính và các rủi ro thuế tiềm ẩn nổi bật của doanh nghiệp.\n\n"
    for name, df in dfs_dict.items():
        if not df.empty:
            prompt += f"--- {name} ---\n"
            prompt += df.to_string(index=False)
            prompt += "\n\n"
    if not risks_df.empty:
        prompt += "--- CÁC RỦI RO ĐÃ PHÁT HIỆN ---\n"
        prompt += risks_df.to_string(index=False)
        prompt += "\n\n"
    if notes_content:
        prompt += "--- NỘI DUNG THUYẾT MINH BCTC ---\n"
        prompt += notes_content[:4000]
        prompt += "\n"

    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(api_url, json=payload, timeout=60)
            response.raise_for_status()
            result = response.json()
            st.session_state['gemini_commentary'] = result['candidates'][0]['content']['parts'][0]['text']
    except Exception as e:
        st.session_state['gemini_commentary'] = f"**Lỗi khi gọi Gemini API:**\n\n{str(e)}"


# --- BÁO CÁO (REPORTING) ---
def display_results():
    st.markdown("---")
    st.header("Kết quả Phân tích & Rà soát")

    tab_titles = ["Tổng hợp & Đối chiếu"]
    if 'notes_content' in st.session_state and st.session_state['notes_content']:
        tab_titles.append("Thuyết minh BCTC")

    if st.session_state.get('output_invoice_data') or st.session_state.get('input_invoice_data'):
        tab_titles.append("Bảng kê Hóa đơn")

    if not st.session_state.get('gtgt_detailed_df', pd.DataFrame()).empty:
        tab_titles.append("Chi tiết 01/GTGT")
    if not st.session_state.get('tndn_main_df', pd.DataFrame()).empty:
        tab_titles.append("Chi tiết 03/TNDN")

    tabs = st.tabs(tab_titles)

    with tabs[0]:
        display_summary_and_risks()

    tab_index = 1
    if 'notes_content' in st.session_state and st.session_state['notes_content']:
        with tabs[tab_index]:
            st.subheader("📝 Nội dung Thuyết minh Báo cáo tài chính")
            st.text_area("Nội dung trích xuất từ file:", st.session_state['notes_content'], height=400)
        tab_index += 1

    if st.session_state.get('output_invoice_data') or st.session_state.get('input_invoice_data'):
        with tabs[tab_index]:
            display_invoice_details()
        tab_index += 1

    if not st.session_state.get('gtgt_detailed_df', pd.DataFrame()).empty:
        with tabs[tab_index]:
            st.subheader("📄 Bảng chi tiết Tờ khai 01/GTGT")
            st.dataframe(st.session_state['gtgt_detailed_df'])
        tab_index += 1

    if not st.session_state.get('tndn_main_df', pd.DataFrame()).empty:
        with tabs[tab_index]:
            st.subheader("📄 Bảng chi tiết Tờ khai chính 03/TNDN")
            st.dataframe(st.session_state['tndn_main_df'])
            st.subheader("📄 Bảng chi tiết Phụ lục 03-1A/TNDN")
            st.dataframe(st.session_state['tndn_appendix_df'])
        tab_index += 1


def display_summary_and_risks():
    gtgt_summary_df = st.session_state.get('gtgt_summary_df', pd.DataFrame())
    balance_sheet_df = st.session_state.get('balance_sheet_df', pd.DataFrame())
    income_statement_df = st.session_state.get('income_statement_df', pd.DataFrame())
    tndn_summary_df = st.session_state.get('tndn_summary_df', pd.DataFrame())
    tncn_kk_summary_df = st.session_state.get('tncn_kk_summary_df', pd.DataFrame())
    tncn_qtt_summary_df = st.session_state.get('tncn_qtt_summary_df', pd.DataFrame())
    output_invoice_data = st.session_state.get('output_invoice_data', None)
    input_invoice_data = st.session_state.get('input_invoice_data', None)
    all_checks = st.session_state.get('all_checks', [])
    api_key = st.session_state.get('gemini_api_key', '')
    notes_content = st.session_state.get('notes_content', None)

    vietnamese_formatter = lambda x: "{:,.2f}".format(x).replace(",", "X").replace(".", ",").replace("X", ".")

    if not gtgt_summary_df.empty:
        st.subheader("📊 Bảng tổng hợp Tờ khai GTGT")
        st.dataframe(
            gtgt_summary_df.style.format(formatter=vietnamese_formatter,
                                         subset=pd.IndexSlice[:, gtgt_summary_df.columns[1:]]))

    if not tndn_summary_df.empty:
        st.subheader("📊 Bảng tổng hợp Quyết toán TNDN")
        st.dataframe(tndn_summary_df.style.format({'Số tiền (VND)': vietnamese_formatter}))

    if not balance_sheet_df.empty:
        st.subheader("📊 Báo cáo Tình hình tài chính (Bảng Cân đối Kế toán)")
        st.dataframe(
            balance_sheet_df.style.format({'Số cuối năm': vietnamese_formatter, 'Số đầu năm': vietnamese_formatter}))

    if not income_statement_df.empty:
        st.subheader("📊 Báo cáo Kết quả Hoạt động Kinh doanh")
        st.dataframe(
            income_statement_df.style.format({'Năm nay': vietnamese_formatter, 'Năm trước': vietnamese_formatter}))

    if not tncn_kk_summary_df.empty:
        st.subheader("📊 Bảng tổng hợp Tờ khai Khấu trừ TNCN (05/KK)")
        st.dataframe(tncn_kk_summary_df.style.format(formatter=vietnamese_formatter,
                                                     subset=pd.IndexSlice[:, tncn_kk_summary_df.columns[1:]]))

    if not tncn_qtt_summary_df.empty:
        st.subheader("📊 Bảng tổng hợp Quyết toán TNCN (05/QTT)")
        st.dataframe(tncn_qtt_summary_df.style.format({'Số liệu': vietnamese_formatter}))

    if output_invoice_data:
        st.subheader("📊 Bảng tổng hợp từ Hóa đơn Bán ra")
        summary_invoice_df = pd.DataFrame.from_dict(output_invoice_data['valid_summary'], orient='index',
                                                    columns=['Số tiền (VND)'])
        summary_invoice_df.index = ['Tổng tiền chưa thuế', 'Tổng tiền thuế', 'Tổng tiền chiết khấu',
                                    'Tổng tiền thanh toán']
        st.dataframe(summary_invoice_df.style.format(vietnamese_formatter))

        if 'mismatch_df' in output_invoice_data and not output_invoice_data['mismatch_df'].empty:
            st.subheader("⚠️ Cảnh báo: Sai lệch Thành tiền trên Bảng kê chi tiết")
            st.warning(
                "Các dòng dưới đây có (Thành tiền) khác với (Số lượng * Đơn giá). Vui lòng kiểm tra lại file gốc.")
            st.dataframe(output_invoice_data['mismatch_df'].style.format(formatter=vietnamese_formatter))

    if input_invoice_data:
        st.subheader("📊 Bảng tổng hợp từ Hóa đơn Mua vào")
        summary_invoice_df = pd.DataFrame.from_dict(input_invoice_data['valid_summary'], orient='index',
                                                    columns=['Số tiền (VND)'])
        summary_invoice_df.index = ['Tổng tiền chưa thuế', 'Tổng tiền thuế', 'Tổng tiền chiết khấu',
                                    'Tổng tiền thanh toán']
        st.dataframe(summary_invoice_df.style.format(vietnamese_formatter))

    st.subheader("🚨 Bảng Đối chiếu & Rà soát Rủi ro")
    df_checks = pd.DataFrame(all_checks)

    def highlight_status(row):
        color = ''
        if row['Trạng thái'] == 'Cảnh báo':
            color = 'background-color: #FFDDDD'
        elif row['Trạng thái'] in ['Khớp', 'OK']:
            color = 'background-color: #D4EDDA'
        elif row['Trạng thái'] == 'Không đủ dữ liệu':
            color = 'background-color: #F0F0F0'
        return [color] * len(row)

    if not df_checks.empty:
        st.dataframe(df_checks.style.apply(highlight_status, axis=1), use_container_width=True)
    else:
        st.info("Chưa thực hiện đối chiếu nào.")

    col1, col2 = st.columns(2)
    with col1:
        dfs_to_export = {
            "TongHop_GTGT": st.session_state.get('gtgt_summary_df', pd.DataFrame()),
            "ChiTiet_01_GTGT": st.session_state.get('gtgt_detailed_df', pd.DataFrame()),
            "TongHop_TNDN": st.session_state.get('tndn_summary_df', pd.DataFrame()),
            "ChiTiet_03_TNDN": st.session_state.get('tndn_main_df', pd.DataFrame()),
            "ChiTiet_PL_03_1A": st.session_state.get('tndn_appendix_df', pd.DataFrame()),
            "BCTHTC_CDKT": st.session_state.get('balance_sheet_df', pd.DataFrame()),
            "BCKQKD": st.session_state.get('income_statement_df', pd.DataFrame()),
            "PL_CDTK": st.session_state.get('trial_balance_df', pd.DataFrame()),
            "TongHop_TNCN_KK": st.session_state.get('tncn_kk_summary_df', pd.DataFrame()),
            "TongHop_TNCN_QTT": st.session_state.get('tncn_qtt_summary_df', pd.DataFrame()),
            "ChiTiet_TNCN_QTT": st.session_state.get('tncn_details_df', pd.DataFrame()),
            "KetQuaDoiChieu": pd.DataFrame(st.session_state.get('all_checks', []))
        }

        if output_invoice_data:
            summary_invoice_df = pd.DataFrame.from_dict(output_invoice_data['valid_summary'], orient='index',
                                                        columns=['Số tiền (VND)'])
            summary_invoice_df.index = ['Tổng tiền chưa thuế', 'Tổng tiền thuế', 'Tổng tiền chiết khấu',
                                        'Tổng tiền thanh toán']
            dfs_to_export["TongHop_HD_DauRa"] = summary_invoice_df
            if 'full_df' in output_invoice_data:
                dfs_to_export["BK_HD_DauRa"] = output_invoice_data['full_df']
            if 'mismatch_df' in output_invoice_data:
                dfs_to_export["HD_DauRa_SaiLech"] = output_invoice_data['mismatch_df']

        if input_invoice_data:
            if 'full_df' in input_invoice_data:
                dfs_to_export["BK_HD_DauVao"] = input_invoice_data['full_df']

        excel_data = convert_df_to_excel(dfs_to_export)
        st.download_button(label="📥 Kết xuất ra Excel", data=excel_data, file_name="Bao_cao_chi_tiet.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    with col2:
        risks_to_report = [r for r in all_checks if r['Trạng thái'] == 'Cảnh báo']
        if st.button("📝 Tạo Thông báo Giải trình (01/KTTT)"):
            if not risks_to_report:
                st.warning("Không có rủi ro nào (Trạng thái 'Cảnh báo') để tạo thông báo.")
            else:
                doc_buffer = create_word_notice(risks_to_report[0])
                st.download_button(label="📥 Tải Thông báo (Word)", data=doc_buffer,
                                   file_name="Thong_bao_giai_trinh.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.header("✨ Nhận xét của Gemini AI")
    if st.button("Phân tích với Gemini"):
        if not api_key:
            st.error("Vui lòng nhập Gemini API Key ở thanh bên trái để sử dụng chức năng này.")
        else:
            with st.spinner("Gemini đang phân tích dữ liệu..."):
                asyncio.run(get_gemini_analysis(api_key, dfs_to_export, pd.DataFrame(risks_to_report), notes_content))
    if 'gemini_commentary' in st.session_state:
        st.markdown(st.session_state['gemini_commentary'])


def display_invoice_details():
    output_invoice_data = st.session_state.get('output_invoice_data', {})
    input_invoice_data = st.session_state.get('input_invoice_data', {})

    vietnamese_formatter = lambda x: "{:,.2f}".format(x).replace(",", "X").replace(".", ",").replace("X", ".")

    def style_invoice_status(row):
        if 'TrangThaiHoaDon' not in row.index:
            return [''] * len(row)

        status = row['TrangThaiHoaDon']
        color = ''
        if status in ['Hóa đơn đã bị thay thế', 'Hóa đơn đã bị xóa bỏ/hủy bỏ']:
            color = 'background-color: #FFC7CE'  # Đỏ nhạt
        elif status in ['Hóa đơn thay thế', 'Hóa đơn điều chỉnh']:
            color = 'background-color: #C6EFCE'  # Xanh nhạt
        elif status == 'Hóa đơn đã bị điều chỉnh':
            color = 'background-color: #FFEB9C'  # Vàng nhạt

        return [color] * len(row)

    if 'full_df' in output_invoice_data and not output_invoice_data['full_df'].empty:
        st.subheader("📋 Bảng kê hóa đơn Bán ra")
        df_to_display = output_invoice_data['full_df']
        styled_df = df_to_display.style.apply(style_invoice_status, axis=1)
        numeric_cols = [col for col in df_to_display.columns if pd.api.types.is_numeric_dtype(df_to_display[col])]
        styled_df = styled_df.format(formatter=vietnamese_formatter, subset=numeric_cols)
        st.dataframe(styled_df)

    if 'full_df' in input_invoice_data and not input_invoice_data['full_df'].empty:
        st.subheader("📋 Bảng kê hóa đơn Mua vào")
        df_to_display = input_invoice_data['full_df']
        styled_df = df_to_display.style.apply(style_invoice_status, axis=1)
        numeric_cols = [col for col in df_to_display.columns if pd.api.types.is_numeric_dtype(df_to_display[col])]
        styled_df = styled_df.format(formatter=vietnamese_formatter, subset=numeric_cols)
        st.dataframe(styled_df)


def convert_df_to_excel(dfs_dict):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        workbook = writer.book
        number_format = workbook.add_format({'num_format': '#,##0'})

        sheet_written = False  # Cờ để theo dõi xem có sheet nào được ghi chưa

        for sheet_name, df in dfs_dict.items():
            # Thêm điều kiện kiểm tra df không phải là None
            if df is not None and not df.empty:
                df.to_excel(writer, index=False, sheet_name=sheet_name)
                worksheet = writer.sheets[sheet_name]
                sheet_written = True  # Đặt cờ thành True khi ghi thành công

                # --- Logic định dạng cột ---
                for idx, col in enumerate(df.columns):
                    series = df[col]
                    # Tính toán độ rộng cột một cách an toàn
                    if not series.empty:
                        max_len = max(
                            series.astype(str).map(len).max(),
                            len(str(series.name))
                        ) + 2
                    else:
                        max_len = len(str(series.name)) + 2

                    worksheet.set_column(idx, idx, max_len)

                    if pd.api.types.is_numeric_dtype(series.dtype) and \
                            "Mã" not in col and "Số hiệu" not in col and "MST" not in col:
                        worksheet.set_column(idx, idx, max_len, number_format)

        # Sau khi duyệt qua tất cả, nếu không có sheet nào được ghi, hãy tạo một sheet mặc định
        if not sheet_written:
            pd.DataFrame({"Thông báo": ["Không có dữ liệu hợp lệ để kết xuất."]}) \
                .to_excel(writer, index=False, sheet_name="Luu_y")

    return output.getvalue()


def create_word_notice(risk_data):
    doc = Document()
    now = datetime.now()
    date_str = f"Quảng Ngãi, ngày {now.day} tháng {now.month} năm {now.year}"
    p = doc.add_paragraph();
    p.add_run('CỤC THUẾ TỈNH QUẢNG NGÃI').bold = True;
    p.alignment = 1
    p = doc.add_paragraph();
    p.add_run(date_str).italic = True;
    p.alignment = 2
    doc.add_heading('THÔNG BÁO', 0)
    doc.add_heading('Về việc giải trình, bổ sung thông tin, tài liệu', level=1)
    doc.add_paragraph(f"Kính gửi: [Tên Người nộp thuế]")
    doc.add_paragraph(f"Mã số thuế: [Mã số thuế]")
    doc.add_paragraph(
        "Căn cứ hồ sơ khai thuế do người nộp thuế nộp tại cơ quan Thuế, Cục Thuế tỉnh Quảng Ngãi đề nghị Quý đơn vị giải trình, cung cấp thông tin, tài liệu sau đây:")
    p = doc.add_paragraph();
    p.add_run('Nội dung cần giải trình: ').bold = True;
    p.add_run(f"{risk_data['Nội dung']} tại kỳ tính thuế {risk_data.get('ky_tinh_thue', 'Cả năm')}.")
    p = doc.add_paragraph();
    p.add_run('Chi tiết: ').bold = True;
    p.add_run(
        f"Đối chiếu {risk_data['Nội dung']}. Số liệu A: {risk_data['Số liệu A']} vs Số liệu B: {risk_data['Số liệu B']}. Chênh lệch: {risk_data.get('Chênh lệch', 'N/A')}. {risk_data.get('Gợi ý', '')}")
    doc.add_paragraph(
        "Văn bản giải trình, cung cấp thông tin, hồ sơ tài liệu gửi về Cục Thuế trong thời hạn 10 ngày làm việc.")
    buffer = io.BytesIO();
    doc.save(buffer);
    buffer.seek(0)
    return buffer


def main():
    params = setup_ui()

    if params["start_button"]:
        for key in list(st.session_state.keys()):
            if key not in ['analysis_complete']:
                del st.session_state[key]

        progress_bar = st.sidebar.progress(0)
        status_text = st.sidebar.empty()
        files_to_analyze = []

        if params["mode"] == "Tự động hóa":
            pass
        else:
            if params.get("uploaded_files"):
                temp_dir = os.path.join(os.getcwd(), "manual_uploads")
                if not os.path.exists(temp_dir): os.makedirs(temp_dir)
                for uploaded_file in params["uploaded_files"]:
                    file_path = os.path.join(temp_dir, uploaded_file.name)
                    with open(file_path, "wb") as f: f.write(uploaded_file.getbuffer())
                    files_to_analyze.append(file_path)
        progress_bar.progress(50)

        status_text.text("Bước 2/5: Xử lý bảng kê hóa đơn...")
        output_invoice_data = None
        if params.get("output_invoice_file"):
            if params["output_invoice_type"] == "Chi tiết":
                output_invoice_data = process_detailed_invoice_data(params["output_invoice_file"])
            else:
                output_invoice_data = process_summary_invoice_data(params["output_invoice_file"])

        input_invoice_data = None
        if params.get("input_invoice_files"):
            input_invoice_data = process_input_invoice_data(params["input_invoice_files"])

        progress_bar.progress(60)

        status_text.text("Bước 3/5: Xử lý Thuyết minh BCTC...")
        notes_content = process_financial_notes(params.get("financial_notes_file"))
        progress_bar.progress(70)

        status_text.text("Bước 4/5: Phân tích dữ liệu...")
        all_declarations, all_checks, gtgt_summary_df, gtgt_detailed_df, balance_sheet_df, income_statement_df, trial_balance_df, tndn_summary_df, tncn_qtt_summary_df, tncn_details_df, tncn_kk_summary_df = parse_and_analyze(
            files_to_analyze, params["accounting_standard"], output_invoice_data, input_invoice_data, notes_content)
        progress_bar.progress(90)

        st.session_state['analysis_complete'] = True
        st.session_state['all_declarations'] = all_declarations
        st.session_state['all_checks'] = all_checks
        st.session_state['gtgt_summary_df'] = gtgt_summary_df
        st.session_state['balance_sheet_df'] = balance_sheet_df
        st.session_state['income_statement_df'] = income_statement_df
        st.session_state['trial_balance_df'] = trial_balance_df
        st.session_state['tndn_summary_df'] = tndn_summary_df
        st.session_state['tncn_qtt_summary_df'] = tncn_qtt_summary_df
        st.session_state['tncn_details_df'] = tncn_details_df
        st.session_state['tncn_kk_summary_df'] = tncn_kk_summary_df
        st.session_state['gemini_api_key'] = params.get("gemini_api_key")
        st.session_state['output_invoice_data'] = output_invoice_data
        st.session_state['input_invoice_data'] = input_invoice_data
        st.session_state['notes_content'] = notes_content
        st.session_state['invoice_analysis_type_at_run'] = params["output_invoice_type"]

        status_text.text("Bước 5/5: Hiển thị kết quả...")
        progress_bar.progress(100)
        status_text.text("Hoàn thành!")
        time.sleep(1)
        st.rerun()

    if st.session_state.get('analysis_complete', False):
        display_results()


if __name__ == "__main__":
    main()
